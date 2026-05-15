from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.chdir('/home/ubuntu/.openclaw/workspace/resume-ai/backend/templates')

def create_classic():
    doc = Document()
    title = doc.add_heading('个人简历', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph('姓名：')
    doc.add_paragraph('电话：')
    doc.add_paragraph('邮箱：')
    doc.add_paragraph('地址：')
    doc.add_heading('教育背景', level=1)
    doc.add_paragraph('时间 | 学校 | 专业 | 学历')
    doc.add_heading('工作经历', level=1)
    doc.add_paragraph('时间 | 公司 | 职位')
    doc.add_paragraph('工作内容描述...')
    doc.add_heading('技能特长', level=1)
    doc.add_paragraph('技能描述...')
    doc.add_heading('自我评价', level=1)
    doc.add_paragraph('自我介绍...')
    doc.save('classic.docx')

def create_modern():
    doc = Document()
    title = doc.add_heading('RESUME', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Profile', level=1)
    doc.add_paragraph('姓名 | 电话 | 邮箱')
    doc.add_heading('Education', level=1)
    doc.add_paragraph('学校 - 专业 - 学历 - 时间')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('公司 - 职位 - 时间')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('技能列表...')
    doc.save('modern.docx')

def create_simple():
    doc = Document()
    doc.add_paragraph('【姓名】')
    doc.add_paragraph('电话： | 邮箱：')
    doc.add_paragraph('')
    doc.add_paragraph('【教育背景】')
    doc.add_paragraph('')
    doc.add_paragraph('【工作经历】')
    doc.add_paragraph('')
    doc.add_paragraph('【技能】')
    doc.save('simple.docx')

def create_creative():
    doc = Document()
    doc.add_heading('个人简历', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('关于我', level=1)
    doc.add_paragraph('姓名： | 电话： | 邮箱：')
    doc.add_heading('教育经历', level=1)
    doc.add_paragraph('')
    doc.add_heading('工作经历', level=1)
    doc.add_paragraph('')
    doc.add_heading('专长技能', level=1)
    doc.add_paragraph('')
    doc.save('creative.docx')

def create_professional():
    doc = Document()
    doc.add_heading('个人简历', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('一、基本信息', level=1)
    doc.add_paragraph('姓名：          电话：          邮箱：')
    doc.add_heading('二、教育背景', level=1)
    doc.add_paragraph('起止时间    学校    专业    学历')
    doc.add_heading('三、工作经历', level=1)
    doc.add_paragraph('起止时间    公司    职位')
    doc.add_heading('四、专业技能', level=1)
    doc.add_paragraph('')
    doc.add_heading('五、自我评价', level=1)
    doc.save('professional.docx')

create_classic()
create_modern()
create_simple()
create_creative()
create_professional()
print('5个模板文件已创建')