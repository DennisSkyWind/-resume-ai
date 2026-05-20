"""
ResumeAI Backend - FastAPIwith User Authentication
简历优化系统后端服务（带用户认证）
"""

from fastapi import FastAPI, HTTPException, Depends, Header, File, UploadFile, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel
from typing import Optional, List
import json
import os
import re
import sqlite3
import hashlib
import secrets
from datetime import datetime, timedelta
import jwt
import random
import tempfile
import logging

# 配置日志 - Vercel兼容（无文件日志）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# API限流配置
RATE_LIMIT_CONFIG = {
    "enabled": True,
    "requests_per_minute": 60,  # 每分钟最多60次请求
    "requests_per_hour": 500,   # 每小时最多500次请求
    "burst_limit": 10,          # 突发请求上限
    "whitelist": ["127.0.0.1", "localhost"]  # 白名单IP
}

# 限流缓存（使用内存缓存）
rate_limit_cache = {}

# 导入邮箱发送模块
from email_sender import send_verification_code, generate_code, get_email_config, EMAIL_CONFIG, send_marketing_email

# 文件解析模块
import pdfplumber
from docx import Document

# PDF生成模块
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO

# 初始化FastAPI
app = FastAPI(
    title="ResumeAI API",
    description="AI简历优化系统 - 为非IT行业求职者提供简历优化服务",
    version="2.0.0"
)

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据目录 - 多环境兼容
# Render: 使用backend目录下的data子目录
# Vercel: 使用/tmp目录（Serverless）
# 本地: 使用项目根目录的data目录
def get_data_dir():
    # 检查环境变量
    if os.environ.get("VERCEL") == "1":
        return "/tmp"
    
    # Render或本地环境
    # 尝试多个可能的路径
    possible_paths = [
        os.path.join(os.path.dirname(__file__), "data"),  # backend/data
        os.path.join(os.path.dirname(__file__), "..", "data"),  # project_root/data
        "/tmp"  # fallback
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
        # 尝试创建目录
        try:
            os.makedirs(path, exist_ok=True)
            return path
        except:
            continue
    
    return "/tmp"

DATA_DIR = get_data_dir()
USER_DB_PATH = os.path.join(DATA_DIR, "users.db")

# 确保数据库和所有表存在（每次启动都检查）
def ensure_tables_exist():
    """确保所有必要的表都存在"""
    conn = sqlite3.connect(USER_DB_PATH)
    conn.row_factory = sqlite3.Row
    
    # 创建所有表（IF NOT EXISTS确保不会重复创建）
    conn.execute('''CREATE TABLE IF NOT EXISTS users
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT UNIQUE,
        password_hash TEXT,
        name TEXT,
        verified INTEGER DEFAULT 0,
        daily_limit INTEGER DEFAULT 5,
        last_reset TEXT,
        last_login TEXT,
        is_admin INTEGER DEFAULT 0,
        is_paid INTEGER DEFAULT 0,
        user_level TEXT DEFAULT 'free',
        level_expires_at TEXT,
        created_at TEXT)''')
    
    # 添加新字段（兼容旧数据）
    cursor = conn.cursor()
    cursor.execute("PRAGMA table_info(users)")
    columns = [col[1] for col in cursor.fetchall()]
    
    if 'user_level' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN user_level TEXT DEFAULT 'free'")
    if 'level_expires_at' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN level_expires_at TEXT")
    
    conn.commit()
    
    conn.execute('''CREATE TABLE IF NOT EXISTS verification_codes
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT,
        code TEXT,
        expires_at TEXT,
        used INTEGER DEFAULT 0,
        created_at TEXT)''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS user_templates
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        template_id TEXT,
        custom_settings TEXT,
        is_default INTEGER DEFAULT 0,
        created_at TEXT,
        updated_at TEXT,
        FOREIGN KEY(user_id) REFERENCES users(id))''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS usage
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        action TEXT,
        industry TEXT,
        created_at TEXT)''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS orders
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        order_id TEXT,
        amount REAL,
        status TEXT,
        created_at TEXT)''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS resumes
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        content TEXT,
        industry TEXT,
        score REAL,
        created_at TEXT)''')
    
    conn.commit()
    
    # 确保字段存在（ALTER TABLE）
    cursor = conn.cursor()
    cursor.execute("PRAGMA table_info(users)")
    columns = [row[1] for row in cursor.fetchall()]
    
    if 'is_admin' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN is_admin INTEGER DEFAULT 0")
    if 'is_paid' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN is_paid INTEGER DEFAULT 0")
    if 'invite_code' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN invite_code TEXT")
    if 'invited_by' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN invited_by INTEGER")
    
    # 创建邀请记录表
    conn.execute('''CREATE TABLE IF NOT EXISTS invitations
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        inviter_id INTEGER,
        invitee_id INTEGER,
        reward_given INTEGER DEFAULT 0,
        created_at TEXT,
        FOREIGN KEY(inviter_id) REFERENCES users(id),
        FOREIGN KEY(invitee_id) REFERENCES users(id))''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS error_logs
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        error_type TEXT,
        error_message TEXT,
        error_stack TEXT,
        user_id INTEGER,
        request_path TEXT,
        request_method TEXT,
        severity TEXT DEFAULT 'error',
        resolved INTEGER DEFAULT 0,
        created_at TEXT)''')
    
    # 创建通知表
    conn.execute('''CREATE TABLE IF NOT EXISTS notifications
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        message TEXT,
        type TEXT DEFAULT 'info',
        related_id INTEGER,
        related_type TEXT,
        is_read INTEGER DEFAULT 0,
        created_at TEXT)''')
    
    # 创建反馈表（增强版）
    conn.execute('''CREATE TABLE IF NOT EXISTS feedback
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        email TEXT,
        type TEXT,
        title TEXT,
        content TEXT,
        priority TEXT DEFAULT 'medium',
        tags TEXT,
        screenshot_url TEXT,
        status TEXT DEFAULT 'pending',
        admin_reply TEXT,
        reply_at TEXT,
        rating INTEGER,
        rating_comment TEXT,
        created_at TEXT)''')
    
    # 确保旧表有新字段（迁移）
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN title TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN priority TEXT DEFAULT 'medium'")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN tags TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN screenshot_url TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN admin_reply TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN reply_at TEXT")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN rating INTEGER")
    except: pass
    try:
        conn.execute("ALTER TABLE feedback ADD COLUMN rating_comment TEXT")
    except: pass
    
    conn.commit()
    
    # 创建用户自定义关键词表（O4-2新增）
    conn.execute('''CREATE TABLE IF NOT EXISTS user_keywords
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        keyword TEXT,
        industry TEXT,
        weight INTEGER DEFAULT 5,
        created_at TEXT)''')
    
    # 初始化管理员账号
    cursor.execute("SELECT * FROM users WHERE email = ?", ('zhwffy@hotmail.com',))
    admin = cursor.fetchone()
    
    if not admin:
        # 创建管理员
        admin_password = "Hiller"
        admin_hash = hashlib.sha256(admin_password.encode()).hexdigest()
        cursor.execute('''INSERT INTO users 
            (email, password_hash, name, verified, is_admin, is_paid, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)''',
            ('zhwffy@hotmail.com', admin_hash, 'Admin', 1, 1, 1, datetime.now().isoformat()))
        conn.commit()
        logger.info("Admin account created: zhwffy@hotmail.com")
    
    conn.close()
    logger.info("Database tables verified and ready")

# 启动时确保表存在
ensure_tables_exist()

# JWT配置（使用固定密钥）
JWT_SECRET = os.environ.get("JWT_SECRET", "resumeai_jwt_secret_key_2026")  # 优先使用环境变量
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 168  # 7天有效期

# 用户等级权益配置
USER_LEVELS = {
    "free": {
        "name": "免费用户",
        "daily_limit": 5,
        "features": ["analyze", "optimize"],
        "price": 0
    },
    "basic": {
        "name": "基础会员",
        "daily_limit": 20,
        "features": ["analyze", "optimize", "upload", "export"],
        "price": 19
    },
    "pro": {
        "name": "专业会员",
        "daily_limit": 50,
        "features": ["analyze", "optimize", "upload", "export", "templates", "history"],
        "price": 49
    },
    "vip": {
        "name": "VIP会员",
        "daily_limit": -1,  # 无限
        "features": ["analyze", "optimize", "upload", "export", "templates", "history", "priority"],
        "price": 99
    }
}

# 免费用户限制
FREE_LIMIT = 5  # 每天免费使用次数

# IP速率限制配置
RATE_LIMIT_WINDOW = 60  # 60秒窗口
RATE_LIMIT_MAX_REQUESTS = 20  # 每分钟最大请求次数
ip_request_counts = {}  # IP请求计数（内存存储）

def check_rate_limit(ip: str) -> bool:
    """检查IP速率限制"""
    now = datetime.now()
    window_start = now - timedelta(seconds=RATE_LIMIT_WINDOW)
    
    # 清理过期记录
    if ip in ip_request_counts:
        ip_request_counts[ip] = [t for t in ip_request_counts[ip] if t > window_start]
    else:
        ip_request_counts[ip] = []
    
    # 检查请求次数
    if len(ip_request_counts[ip]) >= RATE_LIMIT_MAX_REQUESTS:
        return False
    
    # 记录本次请求
    ip_request_counts[ip].append(now)
    return True

# PDF中文字体配置
FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
FONT_BOLD_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
try:
    pdfmetrics.registerFont(TTFont('NotoSansCJK', FONT_PATH, subfontIndex=0))
    pdfmetrics.registerFont(TTFont('NotoSansCJK-Bold', FONT_BOLD_PATH, subfontIndex=0))
    PDF_FONT_NAME = 'NotoSansCJK'
    PDF_FONT_BOLD = 'NotoSansCJK-Bold'
except Exception as e:
    print(f"字体注册失败: {e}")
    PDF_FONT_NAME = 'Helvetica'
    PDF_FONT_BOLD = 'Helvetica-Bold'

# 阿里云Coding API配置
DASHSCOPE_API_KEY = "sk-sp-e8d1076e8dd4461d8d1edf2542f8de68"
DASHSCOPE_BASE_URL = "https://coding.dashscope.aliyuncs.com/v1"
DASHSCOPE_MODEL = "qwen3.5-plus"

# 加载关键词库
def load_keywords():
    """加载关键词数据库"""
    # 首先尝试从backend目录加载
    keywords_file = os.path.join(os.path.dirname(__file__), "keywords.json")
    logger.info(f"尝试加载关键词文件: {keywords_file}")
    
    if os.path.exists(keywords_file):
        try:
            with open(keywords_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.info(f"成功加载关键词数据库，包含 {len(data)} 个行业")
            return data
        except Exception as e:
            logger.error(f"加载关键词文件失败: {e}")
    
    # 兜底：尝试data目录
    keywords_file = os.path.join(DATA_DIR, "keywords.json")
    logger.info(f"尝试从data目录加载: {keywords_file}")
    
    if os.path.exists(keywords_file):
        try:
            with open(keywords_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.info(f"从data目录成功加载关键词数据库")
            return data
        except Exception as e:
            logger.error(f"从data目录加载失败: {e}")
    
    logger.warning("关键词数据库未加载，使用空字典")
    return {}

KEYWORDS_DB = load_keywords()
logger.info(f"KEYWORDS_DB 初始化完成，行业数: {len(KEYWORDS_DB)}")

# ========== API限流检查 ==========

def check_rate_limit(client_ip: str, endpoint: str = "default") -> bool:
    """检查API限流"""
    if not RATE_LIMIT_CONFIG["enabled"]:
        return True
    
    # 白名单IP不受限
    if client_ip in RATE_LIMIT_CONFIG["whitelist"]:
        return True
    
    current_time = datetime.now()
    minute_key = f"{client_ip}:{endpoint}:{current_time.strftime('%Y%m%d%H%M')}"
    hour_key = f"{client_ip}:{endpoint}:{current_time.strftime('%Y%m%d%H')}"
    
    # 检查每分钟限制
    minute_count = rate_limit_cache.get(minute_key, 0)
    if minute_count >= RATE_LIMIT_CONFIG["requests_per_minute"]:
        logger.warning(f"Rate limit exceeded for {client_ip}: {minute_count} requests/minute")
        return False
    
    # 检查每小时限制
    hour_count = rate_limit_cache.get(hour_key, 0)
    if hour_count >= RATE_LIMIT_CONFIG["requests_per_hour"]:
        logger.warning(f"Rate limit exceeded for {client_ip}: {hour_count} requests/hour")
        return False
    
    # 更新计数
    rate_limit_cache[minute_key] = minute_count + 1
    rate_limit_cache[hour_key] = hour_count + 1
    
    return True

def get_client_ip(request) -> str:
    """获取客户端IP"""
    # 尝试从headers获取真实IP（适用于代理环境）
    forwarded = request.headers.get("X-Forwarded-For")
    if forwarded:
        return forwarded.split(",")[0].strip()
    
    real_ip = request.headers.get("X-Real-IP")
    if real_ip:
        return real_ip
    
    # 直接连接的IP
    if hasattr(request, "client") and request.client:
        return request.client.host
    
    return "unknown"

from fastapi import Request

@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    """限流中间件"""
    # 只对API路径进行限流
    if not request.url.path.startswith("/api"):
        return await call_next(request)
    
    client_ip = get_client_ip(request)
    
    if not check_rate_limit(client_ip, request.url.path):
        return JSONResponse(
            status_code=429,
            content={
                "success": False,
                "detail": "请求过于频繁，请稍后再试",
                "retry_after": 60
            }
        )
    
    response = await call_next(request)
    return response

# ========== 用户认证函数 ==========

def get_user_db():
    conn = sqlite3.connect(USER_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def generate_invite_code() -> str:
    """生成6位邀请码"""
    import random
    import string
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))

def check_password_strength(password: str) -> dict:
    """检查密码强度"""
    result = {
        "valid": True,
        "score": 0,
        "message": ""
    }
    
    # 长度检查
    if len(password) < 6:
        result["valid"] = False
        result["message"] = "密码长度至少6位"
        return result
    elif len(password) >= 8:
        result["score"] += 1
    
    # 复杂度检查
    has_lower = any(c.islower() for c in password)
    has_upper = any(c.isupper() for c in password)
    has_digit = any(c.isdigit() for c in password)
    has_special = any(c in "!@#$%^&*()_+-=[]{}|;:,.<>?/" for c in password)
    
    complexity = sum([has_lower, has_upper, has_digit, has_special])
    result["score"] += complexity
    
    # 评分
    if result["score"] >= 4:
        result["message"] = "密码强度：强"
    elif result["score"] >= 2:
        result["message"] = "密码强度：中等"
    else:
        result["message"] = "密码强度：弱，建议包含大小写字母、数字和特殊字符"
    
    return result

def create_token(user_id: int, email: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def verify_token(token: str) -> Optional[dict]:
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def detect_language(text: str) -> str:
    """检测文本语言（中文/英文）"""
    # 统计中文字符比例
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(text.strip())
    
    if total_chars == 0:
        return "zh"  # 默认中文
    
    chinese_ratio = chinese_chars / total_chars
    
    # 中文字符占比超过30%视为中文
    if chinese_ratio > 0.3:
        return "zh"
    else:
        return "en"

async def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="未提供认证token")
    
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    payload = verify_token(token)
    
    if not payload:
        raise HTTPException(status_code=401, detail="token无效或已过期")
    
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (payload["user_id"],))
    user = cursor.fetchone()
    conn.close()
    
    if not user:
        raise HTTPException(status_code=401, detail="用户不存在")
    
    return dict(user)

async def get_current_user_optional(authorization: Optional[str] = Header(None)):
    """可选的用户认证（不强制要求登录）"""
    if not authorization:
        return None
    
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    payload = verify_token(token)
    
    if not payload:
        return None
    
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (payload["user_id"],))
    user = cursor.fetchone()
    conn.close()
    
    if not user:
        return None
    
    return dict(user)

def get_user_level_info(user_id: int) -> dict:
    """获取用户等级信息"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT user_level, level_expires_at FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()
    conn.close()
    
    level = user["user_level"] if user else "free"
    
    # 检查等级是否过期
    if user and user["level_expires_at"]:
        expires = datetime.fromisoformat(user["level_expires_at"])
        if datetime.now() > expires:
            level = "free"  # 过期降级
    
    return USER_LEVELS.get(level, USER_LEVELS["free"])

def check_usage_limit(user_id: int) -> bool:
    """检查用户今日使用次数（根据等级）"""
    level_info = get_user_level_info(user_id)
    
    # VIP用户无限制
    if level_info["daily_limit"] == -1:
        return True
    
    # 检查今日使用次数
    conn = get_user_db()
    cursor = conn.cursor()
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user_id, today))
    
    result = cursor.fetchone()
    conn.close()
    
    return result["count"] < level_info["daily_limit"]

def record_usage(user_id: int, action: str, industry: str = None):
    """记录用户使用"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO usage (user_id, action, industry)
        VALUES (?, ?, ?)
    """, (user_id, action, industry))
    conn.commit()
    conn.close()

# ========== 请求模型 ==========

class SendCodeRequest(BaseModel):
    email: str

class RegisterRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = None
    code: str  # 验证码
    invited_by: Optional[str] = None  # 邀请码

class LoginRequest(BaseModel):
    email: str
    password: str

class FeedbackRequest(BaseModel):
    type: str
    title: Optional[str] = None  # 新增：反馈标题

# O4-2新增：自定义关键词请求类
class AddKeywordRequest(BaseModel):
    keyword: str
    industry: str
    weight: int = 5

class UserKeywordsRequest(BaseModel):
    industry: str
    content: str
    priority: Optional[str] = "medium"  # 新增：优先级 low/medium/high/urgent
    tags: Optional[List[str]] = None  # 新增：标签列表
    screenshot: Optional[str] = None  # 新增：截图URL（base64或URL）
    email: Optional[str] = None
    page: Optional[str] = None

class AnalyzeRequest(BaseModel):
    resume_content: str
    industry: str
    sub_industry: Optional[str] = None  # 新增O4-1：子行业参数
    language: str = "zh"
    jd_content: Optional[str] = None
    # O5-1新增：评分权重配置
    keyword_weight: Optional[int] = 60  # 关键词权重百分比（40-80）
    ats_weight: Optional[int] = 40  # ATS权重百分比（20-60）

class OptimizeRequest(BaseModel):
    resume_content: str
    industry: str
    sub_industry: Optional[str] = None  # 新增O4-1：子行业参数
    language: str = "zh"
    jd_content: Optional[str] = None

# ========== 用户认证API ==========

# 显式处理OPTIONS预检请求（确保CORS正常工作）
@app.options("/api/v1/{path:path}")
async def options_handler():
    return Response(status_code=200)

@app.post("/api/v1/send-code")
async def send_code(request: SendCodeRequest):
    """发送验证码"""
    try:
        import sqlite3
        
        conn = get_user_db()
        cursor = conn.cursor()
        
        # 检查邮箱是否已注册
        cursor.execute("SELECT id FROM users WHERE email = ?", (request.email,))
        if cursor.fetchone():
            conn.close()
            raise HTTPException(status_code=400, detail="邮箱已被注册")
        
        # 生成验证码
        code = generate_code(EMAIL_CONFIG["code_length"])
        expires_at = datetime.now() + timedelta(minutes=EMAIL_CONFIG["code_expire_minutes"])
        
        # 保存验证码到数据库
        cursor.execute("""
            INSERT INTO verification_codes (email, code, expires_at)
            VALUES (?, ?, ?)
        """, (request.email, code, expires_at.isoformat()))
        conn.commit()
        conn.close()
        
        # 发送验证码
        result = send_verification_code(request.email, code)
        
        return {
            "success": result["success"],
            "message": result["message"],
            "test_code": result.get("test_code")  # 测试模式返回验证码
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"发送验证码失败: {e}")
        raise HTTPException(status_code=500, detail=f"发送失败: {str(e)}")

@app.post("/api/v1/register")
async def register(request: RegisterRequest):
    """用户注册（需要验证码）"""
    logger.info(f"注册请求: email={request.email}")
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 验证邮箱格式
    if not re.match(r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$", request.email):
        logger.warning(f"注册失败: {request.email} - 邮箱格式不正确")
        raise HTTPException(status_code=400, detail="邮箱格式不正确")
    
    # 检查密码强度
    strength = check_password_strength(request.password)
    if not strength["valid"]:
        logger.warning(f"注册失败: {request.email} - 密码强度不足")
        raise HTTPException(status_code=400, detail=strength["message"])
    
    # 检查邮箱是否已存在
    cursor.execute("SELECT id FROM users WHERE email = ?", (request.email,))
    if cursor.fetchone():
        conn.close()
        logger.warning(f"注册失败: {request.email} - 邮箱已被注册")
        raise HTTPException(status_code=400, detail="邮箱已被注册")
    
    # 验证验证码
    cursor.execute("""
        SELECT * FROM verification_codes 
        WHERE email = ? AND code = ? AND used = FALSE AND expires_at > ?
        ORDER BY created_at DESC LIMIT 1
    """, (request.email, request.code, datetime.now().isoformat()))
    
    code_record = cursor.fetchone()
    
    if not code_record:
        conn.close()
        raise HTTPException(status_code=400, detail="验证码无效或已过期")
    
    # 标记验证码已使用
    cursor.execute("UPDATE verification_codes SET used = TRUE WHERE id = ?", (code_record["id"],))
    
    # 创建用户
    password_hash = hash_password(request.password)
    
    # 生成用户邀请码（6位字母数字）
    user_invite_code = generate_invite_code()
    
    # 处理邀请关系
    inviter_id = None
    if request.invited_by:
        cursor.execute("SELECT id FROM users WHERE invite_code = ?", (request.invited_by,))
        inviter = cursor.fetchone()
        if inviter:
            inviter_id = inviter["id"]
    
    cursor.execute("""
        INSERT INTO users (email, password_hash, name, invite_code, invited_by)
        VALUES (?, ?, ?, ?, ?)
    """, (request.email, password_hash, request.name, user_invite_code, inviter_id))
    
    user_id = cursor.lastrowid
    
    # 如果有邀请人，记录邀请关系并奖励
    if inviter_id:
        cursor.execute("""
            INSERT INTO invitations (inviter_id, invitee_id, reward_given, created_at)
            VALUES (?, ?, ?, ?)
        """, (inviter_id, user_id, 0, datetime.now().isoformat()))
        
        # 给邀请人增加3次使用额度
        cursor.execute("UPDATE users SET daily_limit = daily_limit + 3 WHERE id = ?", (inviter_id,))
        # 给新用户增加1次使用额度
        cursor.execute("UPDATE users SET daily_limit = daily_limit + 1 WHERE id = ?", (user_id,))
    
    conn.commit()
    conn.close()
    
    # 生成token
    token = create_token(user_id, request.email)
    
    return {
        "success": True,
        "data": {
            "user_id": user_id,
            "email": request.email,
            "token": token,
            "token_type": "Bearer"
        }
    }

@app.post("/api/v1/login")
async def login(request: LoginRequest):
    """用户登录"""
    logger.info(f"登录请求: {request.email}")
    conn = get_user_db()
    cursor = conn.cursor()
    
    password_hash = hash_password(request.password)
    cursor.execute("""
            SELECT id, email, name, daily_limit, is_admin, is_paid, user_level FROM users WHERE email = ? AND password_hash = ?
        """, (request.email, password_hash))
    
    user = cursor.fetchone()
    
    if not user:
        logger.warning(f"登录失败: {request.email} - 邮箱或密码错误")
        conn.close()
        raise HTTPException(status_code=401, detail="邮箱或密码错误")
    
    # 更新登录时间
    cursor.execute("""
        UPDATE users SET last_login = ? WHERE id = ?
    """, (datetime.now().isoformat(), user["id"]))
    conn.commit()
    conn.close()
    
    # 生成token
    token = create_token(user["id"], user["email"])
    
    return {
        "success": True,
        "data": {
            "user_id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "daily_limit": user["daily_limit"],
            "is_admin": bool(user["is_admin"] if "is_admin" in user.keys() else 0),
            "is_paid": user["is_paid"] if "is_paid" in user.keys() else 0,
            "user_level": user["user_level"] if "user_level" in user.keys() else "free",
            "token": token,
            "token_type": "Bearer"
        }
    }

@app.get("/api/v1/me")
async def get_me(user: dict = Depends(get_current_user)):
    """获取当前用户信息（包含等级详情）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取用户等级信息
    user_level = user["user_level"] if "user_level" in user.keys() else "free"
    level_expires_at = user["level_expires_at"] if "level_expires_at" in user.keys() else None
    
    # 检查等级是否过期
    is_expired = False
    if level_expires_at:
        try:
            expires = datetime.fromisoformat(level_expires_at)
            if datetime.now() > expires:
                is_expired = True
                user_level = "free"  # 过期降级
        except:
            pass
    
    # 获取等级权益
    level_info = USER_LEVELS.get(user_level, USER_LEVELS["free"])
    daily_limit = level_info["daily_limit"]
    
    # 获取今日使用次数
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user["id"], today))
    
    usage_today = cursor.fetchone()["count"]
    
    # 计算剩余次数
    if daily_limit == -1:
        remaining = "无限"
        usage_percentage = 0
    else:
        remaining = max(0, daily_limit - usage_today)
        usage_percentage = min(100, int(usage_today / daily_limit * 100))
    
    # 获取简历数量
    cursor.execute("SELECT COUNT(*) as count FROM resumes WHERE user_id = ?", (user["id"],))
    resumes_count = cursor.fetchone()["count"]
    conn.close()
    
    return {
        "success": True,
        "data": {
            "user_id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "user_level": user_level,
            "level_name": level_info["name"],
            "level_features": level_info["features"],
            "daily_limit": daily_limit,
            "usage_today": usage_today,
            "remaining_usage": remaining,
            "usage_percentage": usage_percentage,
            "level_expires_at": level_expires_at,
            "is_expired": is_expired,
            "is_admin": bool(user["is_admin"] if "is_admin" in user.keys() else 0),
            "resumes_count": resumes_count
        }
    }

@app.get("/api/v1/usage")
async def get_usage(user: dict = Depends(get_current_user)):
    """获取用户使用次数详情"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取今日使用次数
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user["id"], today))
    
    usage = cursor.fetchone()
    used_today = usage["count"]
    
    # 获取是否付费
    is_paid = user["is_paid"]
    remaining = "unlimited" if is_paid else max(0, FREE_LIMIT - used_today)
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "used_today": used_today,
            "limit": FREE_LIMIT if not is_paid else "unlimited",
            "remaining": remaining,
            "is_paid": is_paid,
            "can_use": is_paid or used_today < FREE_LIMIT
        }
    }

@app.get("/api/v1/invite/code")
async def get_invite_code(user: dict = Depends(get_current_user)):
    """获取用户邀请码"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取用户邀请码
    cursor.execute("SELECT invite_code FROM users WHERE id = ?", (user["id"],))
    row = cursor.fetchone()
    
    invite_code = row["invite_code"] if row and row["invite_code"] else ""
    
    # 如果没有邀请码，生成一个
    if not invite_code:
        invite_code = generate_invite_code()
        cursor.execute("UPDATE users SET invite_code = ? WHERE id = ?", (invite_code, user["id"]))
        conn.commit()
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "invite_code": invite_code,
            "invite_link": f"https://resume.bearai.bond?invite={invite_code}"
        }
    }

@app.get("/api/v1/invite/stats")
async def get_invite_stats(user: dict = Depends(get_current_user)):
    """获取邀请统计"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取邀请人数
    cursor.execute("""
        SELECT COUNT(*) as count FROM invitations WHERE inviter_id = ?
    """, (user["id"],))
    invite_count = cursor.fetchone()["count"]
    
    # 获取已奖励次数
    cursor.execute("""
        SELECT SUM(reward_given) as total FROM invitations WHERE inviter_id = ?
    """, (user["id"],))
    total_reward = cursor.fetchone()["total"] or 0
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "invite_count": invite_count,
            "total_reward": total_reward * 3,  # 每次邀请奖励3次
            "reward_per_invite": 3
        }
    }

# ========== 简历API（需认证） ==========

@app.post("/api/v1/upload")
async def upload_resume(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """上传简历文件并提取文本（PDF/Word）"""
    
    # 检查文件类型
    filename = file.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.docx') or filename.endswith('.doc')):
        raise HTTPException(status_code=400, detail="只支持PDF和Word文件（.pdf, .docx, .doc）")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完")
    
    try:
        # 保存临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        text_content = ""
        
        # 解析PDF
        if filename.endswith('.pdf'):
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        
        # 解析Word
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text_content += " | ".join(row_text) + "\n"
        
        # 清理临时文件
        os.unlink(tmp_path)
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="无法从文件中提取文本内容")
        
        # 记录使用
        record_usage(user["id"], "upload", "file")
        
        return {
            "success": True,
            "data": {
                "filename": file.filename,
                "content": text_content.strip(),
                "length": len(text_content.strip())
            }
        }
        
    except Exception as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

@app.post("/api/v1/analyze")
async def analyze(request: AnalyzeRequest, user: dict = Depends(get_current_user), 
                  client_ip: Optional[str] = Header(None, alias="X-Forwarded-For")):
    """分析简历（需认证）"""
    # IP速率限制
    ip = client_ip or "unknown"
    if not check_rate_limit(ip):
        logger.warning(f"速率限制触发: IP={ip}")
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    
    logger.info(f"简历分析请求: user={user['id']}, industry={request.industry}, ip={ip}")
    
    if request.industry not in KEYWORDS_DB:
        raise HTTPException(status_code=400, detail=f"不支持的行业: {request.industry}")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完，请升级付费或明天再试")
    
    # O4-2新增：获取用户自定义关键词
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT keyword, weight FROM user_keywords WHERE user_id = ? AND industry = ?
        ORDER BY weight DESC
    """, (user["id"], request.industry))
    user_keywords_list = [{"keyword": row[0], "weight": row[1]} for row in cursor.fetchall()]
    
    # 分析简历（O4-1/O4-2新增参数）
    result = analyze_resume(request.resume_content, request.industry, request.language, request.jd_content, request.sub_industry, user_keywords_list, request.keyword_weight, request.ats_weight)
    
    # 记录使用
    record_usage(user["id"], "analyze", request.industry)
    
    # 保存简历
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO resumes (user_id, content, industry, score)
        VALUES (?, ?, ?, ?)
    """, (user["id"], request.resume_content, request.industry, result["overall_score"]))
    conn.commit()
    conn.close()
    
    return {"success": True, "data": result}

@app.post("/api/v1/optimize")
async def optimize(request: OptimizeRequest, user: dict = Depends(get_current_user)):
    """优化简历（需认证）"""
    if request.industry not in KEYWORDS_DB:
        raise HTTPException(status_code=400, detail=f"不支持的行业: {request.industry}")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完")
    
    # 优化简历（简化版）
    result = optimize_resume(request.resume_content, request.industry, request.language)
    
    # 记录使用
    record_usage(user["id"], "optimize", request.industry)
    
    return {"success": True, "data": result}

@app.get("/api/v1/resumes")
async def get_resumes(user: dict = Depends(get_current_user)):
    """获取用户简历历史"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, title, industry, score, created_at
        FROM resumes WHERE user_id = ?
        ORDER BY created_at DESC LIMIT 20
    """, (user["id"],))
    
    resumes = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return {"success": True, "data": resumes}

@app.get("/api/v1/resumes/history")
async def get_resumes_history(user: dict = Depends(get_current_user)):
    """获取用户简历历史列表（带分页）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取总数
    cursor.execute("SELECT COUNT(*) as total FROM resumes WHERE user_id = ?", (user["id"],))
    total = cursor.fetchone()["total"]
    
    # 获取历史列表（带更多字段）
    cursor.execute("""
        SELECT 
            id,
            title,
            industry,
            score,
            created_at,
            updated_at,
            LENGTH(content) as content_length
        FROM resumes 
        WHERE user_id = ?
        ORDER BY created_at DESC
        LIMIT 50
    """, (user["id"],))
    
    history = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return {
        "success": True,
        "data": {
            "total": total,
            "history": history
        }
    }

@app.get("/api/v1/resumes/{resume_id}")
async def get_resume_detail(resume_id: int, user: dict = Depends(get_current_user)):
    """获取简历详情"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            id,
            user_id,
            title,
            industry,
            score,
            content,
            created_at,
            updated_at
        FROM resumes 
        WHERE id = ? AND user_id = ?
    """, (resume_id, user["id"]))
    
    resume = cursor.fetchone()
    conn.close()
    
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在或无权访问")
    
    return {
        "success": True,
        "data": dict(resume)
    }

@app.post("/api/v1/export/pdf")
async def export_pdf(user: dict = Depends(get_current_user)):
    """导出简历分析报告PDF"""
    # 生成PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    
    # 样式配置
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        fontName=PDF_FONT_BOLD,
        fontSize=18,
        spaceAfter=20,
        alignment=1  # 居中
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        fontName=PDF_FONT_BOLD,
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'CustomBody',
        fontName=PDF_FONT_NAME,
        fontSize=12,
        spaceBefore=6,
        spaceAfter=6,
        leading=18
    )
    
    # 构建PDF内容
    elements = []
    
    # 标题
    elements.append(Paragraph("简历分析报告", title_style))
    elements.append(Spacer(1, 20))
    
    # 基本信息
    elements.append(Paragraph("基本信息", heading_style))
    info_data = [
        ["分析日期", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["用户邮箱", user.get("email", "未知")],
    ]
    info_table = Table(info_data, colWidths=[80*mm, 80*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), PDF_FONT_NAME),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 30))
    
    # 落款
    elements.append(Spacer(1, 50))
    elements.append(Paragraph(
        f"{datetime.now().strftime('%Y年%m月%d日')} 智能与数字化中心·数据智能",
        ParagraphStyle('Footer', fontName=PDF_FONT_NAME, fontSize=12, alignment=2)
    ))
    
    # 生成PDF
    doc.build(elements)
    buffer.seek(0)
    
    # 返回PDF响应
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=resume_analysis_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

@app.post("/api/v1/export/word")
async def export_word(user: dict = Depends(get_current_user)):
    """导出简历优化结果为Word文档"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    
    # 创建Word文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('简历优化报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph(f'分析日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'用户邮箱: {user.get("email", "未知")}')
    
    # 说明
    doc.add_heading('使用说明', level=1)
    doc.add_paragraph('本报告由 ResumeAI 自动生成，供求职者参考使用。')
    doc.add_paragraph('建议根据分析结果优化简历内容，提高求职成功率。')
    
    # 落款
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph(f'{datetime.now().strftime("%Y年%m月%d日")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer2 = doc.add_paragraph('智能与数字化中心·数据智能')
    footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=resume_optimized_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

# ========== 核心分析函数 ==========

def analyze_resume(resume: str, industry: str, language: str, jd: Optional[str] = None, sub_industry: Optional[str] = None, user_keywords: Optional[list] = None, keyword_weight_pct: int = 60, ats_weight_pct: int = 40):
    """分析简历关键词匹配度（新增sub_industry、user_keywords、权重参数O4-1/O4-2/O5-1）"""
    
    industry_keywords = KEYWORDS_DB.get(industry, {})
    required_keywords = industry_keywords.get("keywords", {}).get("required", [])
    preferred_keywords = industry_keywords.get("keywords", {}).get("preferred", [])
    
    # 如果有子行业，添加子行业专属关键词（O4-1）
    sub_industry_keywords = []
    if sub_industry and "sub_industries" in industry_keywords:
        for sub in industry_keywords["sub_industries"]:
            if sub.get("id") == sub_industry:
                sub_industry_keywords = sub.get("keywords_boost", [])
                # 将子行业关键词添加为preferred关键词（权重+2）
                for kw in sub_industry_keywords:
                    preferred_keywords.append({
                        "keyword": kw,
                        "weight": 7,  # 子行业关键词中等权重
                        "source": "sub_industry"
                    })
                break
    
    # O4-2新增：添加用户自定义关键词
    if user_keywords:
        for kw in user_keywords:
            preferred_keywords.append({
                "keyword": kw.get("keyword"),
                "weight": kw.get("weight", 8),  # 用户关键词默认高权重
                "source": "user_custom"
            })
    
    all_keywords = required_keywords + preferred_keywords
    
    keywords_found = []
    keywords_missing = []
    total_weight = 0
    matched_weight = 0
    
    for kw in all_keywords:
        keyword_text = kw["keyword"] if language == "zh" else kw.get("en", kw["keyword"])
        count = len(re.findall(keyword_text, resume, re.IGNORECASE))
        found = count > 0
        
        if found:
            keywords_found.append({
                "keyword": kw["keyword"],
                "weight": kw["weight"],
                "count": count,
                "found": True
            })
            matched_weight += kw["weight"]
        
        total_weight += kw["weight"]
    
    # JD匹配度分析
    jd_match_result = None
    if jd:
        jd_keywords = extract_jd_keywords(jd)
        jd_matched = []
        jd_missing = []
        jd_total = len(jd_keywords)
        jd_matched_count = 0
        
        for kw in jd_keywords:
            if re.search(kw, resume, re.IGNORECASE):
                jd_matched.append(kw)
                jd_matched_count += 1
            else:
                jd_missing.append(kw)
        
        jd_match_score = round(jd_matched_count / jd_total * 100, 1) if jd_total > 0 else 0
        jd_match_result = {
            "jd_keywords": jd_keywords,
            "jd_matched": jd_matched,
            "jd_missing": jd_missing,
            "jd_match_score": jd_match_score,
            "jd_total": jd_total,
            "jd_matched_count": jd_matched_count
        }
    
    # 计算缺失关键词
    found_kw_list = [k["keyword"] for k in keywords_found]
    for kw in all_keywords:
        if kw["keyword"] not in found_kw_list:
            keywords_missing.append({
                "keyword": kw["keyword"],
                "weight": kw["weight"],
                "suggestion": f"添加{kw['keyword']}相关经验描述"
            })
    
    # 计算评分（包含JD匹配度）
    keyword_score = round(matched_weight / total_weight * 100, 1) if total_weight > 0 else 0
    ats_result = check_ats(resume)
    ats_score = ats_result["score"]
    
    # 关键词密度分析（新增）
    keyword_density_result = analyze_keyword_density(resume, keywords_found, all_keywords)
    
    # 量化指标分析（新增）
    quantified_metrics_result = analyze_quantified_metrics(resume)
    
    # O5-1新增：使用用户自定义权重计算综合评分
    # 将百分比转换为权重系数（keyword_weight_pct默认60，ats_weight_pct默认40）
    kw_weight = keyword_weight_pct / 100.0  # 转换为0-1范围
    ats_weight = ats_weight_pct / 100.0
    
    # 如果有JD，综合评分加入JD匹配度
    if jd_match_result:
        # 有JD时，关键词权重降低，JD权重提高
        jd_weight = 0.5
        adjusted_kw_weight = kw_weight * 0.6  # 调整为30%
        adjusted_ats_weight = ats_weight * 0.4  # 调整为20%
        overall_score = round((keyword_score * adjusted_kw_weight + ats_score * adjusted_ats_weight + jd_match_result["jd_match_score"] * jd_weight), 1)
    else:
        # 无JD时，按用户权重计算
        overall_score = round((keyword_score * kw_weight + ats_score * ats_weight), 1)
    
    suggestions = generate_suggestions(resume, keywords_missing, industry)
    
    # 添加量化指标建议
    if quantified_metrics_result["suggestions"]:
        suggestions.extend(quantified_metrics_result["suggestions"])
    
    # 如果有JD缺失关键词，添加到建议
    if jd_match_result and jd_match_result["jd_missing"]:
        for kw in jd_match_result["jd_missing"][:5]:  # 最多显示5个
            suggestions.append({
                "type": "jd_missing",
                "keyword": kw,
                "suggestion": f"根据职位要求，建议添加'{kw}'相关经验"
            })
    
    # 为缺失关键词生成示例句子（新增）
    example_sentences = {}
    for kw in keywords_missing[:5]:
        kw_text = kw["keyword"]
        examples = generate_example_sentences(kw_text, industry)
        example_sentences[kw_text] = examples
    
    # 为缺失关键词推荐位置（新增O3-2）
    position_suggestions = {}
    if keywords_missing:
        try:
            position_result = suggest_keyword_positions(resume, keywords_missing[:8], industry)
            for suggestion in position_result.get("suggestions", []):
                position_suggestions[suggestion["keyword"]] = {
                    "recommended_section": suggestion["recommended_section"],
                    "alternatives": suggestion.get("alternative_sections", []),
                    "suggestion_text": suggestion["suggestion_text"]
                }
        except Exception as e:
            logger.warning(f"位置建议生成失败: {e}")
    
    # 同义词推荐（新增O3-3）
    synonym_suggestions = {}
    for kw in keywords_found[:10]:  # 已找到的关键词
        kw_text = kw["keyword"]
        synonyms = get_keyword_synonyms(kw_text, industry)
        if synonyms:
            synonym_suggestions[kw_text] = synonyms
    
    # 为缺失关键词也提供同义词
    for kw in keywords_missing[:5]:
        kw_text = kw["keyword"]
        synonyms = get_keyword_synonyms(kw_text, industry)
        if synonyms:
            synonym_suggestions[kw_text] = synonyms
    
    return {
        "overall_score": overall_score,
        "keyword_score": keyword_score,
        "ats_score": ats_score,
        "ats_details": ats_result,
        "keyword_density": keyword_density_result,
        "metrics_analysis": quantified_metrics_result,
        "jd_match": jd_match_result,
        "keywords_found": keywords_found,
        "keywords_missing": keywords_missing,
        "suggestions": suggestions,
        "example_sentences": example_sentences,  # 新增：示例句子
        "position_suggestions": position_suggestions,  # 新增O3-2：位置建议
        "synonym_suggestions": synonym_suggestions,  # 新增O3-3：同义词推荐
        "industry_name": industry_keywords.get("name", industry)
    }

def extract_jd_keywords(jd: str) -> list:
    """从JD中提取关键词"""
    # 常见技能关键词模式
    skill_patterns = [
        r'[a-zA-Z]+(?:\+[+]|#|\.js|\.py|\.ts)?',  # 技术名词如 Python, C++, React.js
        r'(?:熟悉|精通|掌握|了解|能够|熟练)[\s的]*([^，。；\n]{2,8})',  # 中文技能描述
        r'(\d+)\s*[年月天]',  # 经验年限
    ]
    
    keywords = set()
    
    # 提取英文技能词
    en_skills = re.findall(r'([A-Z][a-zA-Z]*(?:\+[+]|#|\.js|\.py|\.ts)?|[a-z]+(?:\+[+]|#))', jd)
    keywords.update(en_skills)
    
    # 提取常见技能关键词
    common_skills = [
        "Python", "Java", "JavaScript", "C++", "Go", "Rust", "SQL", "MySQL", "PostgreSQL",
        "React", "Vue", "Angular", "Node", "Django", "Spring", "Docker", "Kubernetes",
        "AWS", "Azure", "GCP", "Linux", "Git", "API", "REST", "GraphQL",
        "沟通", "协作", "分析", "设计", "开发", "测试", "运维", "管理",
        "Excel", "PowerPoint", "Word", "Office", "数据分析", "项目管理"
    ]
    
    for skill in common_skills:
        if skill.lower() in jd.lower() or skill in jd:
            keywords.add(skill)
    
    return list(keywords)[:20]  # 最多返回20个关键词

def optimize_resume(resume: str, industry: str, language: str):
    """优化简历"""
    analysis = analyze_resume(resume, industry, language)
    
    optimized = resume
    
    # 添加建议
    suggestions_text = "\n\n".join([s["suggestion"] for s in analysis["suggestions"]])
    
    return {
        "optimized_resume": optimized,
        "added_suggestions": suggestions_text,
        "score_improvement": f"建议添加{len(analysis['keywords_missing'])}个关键词",
        "analysis": analysis
    }

def check_ats(resume: str) -> dict:
    """ATS兼容性检测 - 扩展版（12项检测）"""
    checks = []
    total_score = 100
    
    # 1. 简历长度检测
    length = len(resume)
    if length < 200:
        checks.append({"name": "简历长度", "status": "warning", "message": f"简历过短（{length}字），建议至少300字以上", "score": 80})
        total_score -= 20
    elif length > 5000:
        checks.append({"name": "简历长度", "status": "warning", "message": f"简历过长（{length}字），建议控制在5000字以内", "score": 90})
        total_score -= 10
    else:
        checks.append({"name": "简历长度", "status": "pass", "message": f"简历长度适中（{length}字）", "score": 100})
    
    # 2. 电话号码检测
    phone_pattern = r'(1[3-9]\d{9}|0\d{2,3}-?\d{7,8}|\d{3,4}-?\d{3,4}-?\d{4})'
    has_phone = re.search(phone_pattern, resume)
    if has_phone:
        checks.append({"name": "电话号码", "status": "pass", "message": "检测到电话号码", "score": 100})
    else:
        checks.append({"name": "电话号码", "status": "error", "message": "缺少电话号码", "score": 0})
        total_score -= 15
    
    # 3. 邮箱检测
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    has_email = re.search(email_pattern, resume)
    if has_email:
        checks.append({"name": "电子邮箱", "status": "pass", "message": "检测到邮箱地址", "score": 100})
    else:
        checks.append({"name": "电子邮箱", "status": "error", "message": "缺少邮箱地址", "score": 0})
        total_score -= 15
    
    # 4. 段落结构检测（按换行符分割）
    paragraphs = [p.strip() for p in resume.split('\n\n') if p.strip()]
    if len(paragraphs) >= 3:
        checks.append({"name": "段落结构", "status": "pass", "message": f"有{len(paragraphs)}个清晰段落", "score": 100})
    elif len(paragraphs) >= 2:
        checks.append({"name": "段落结构", "status": "warning", "message": f"仅{len(paragraphs)}个段落，建议增加段落划分", "score": 70})
        total_score -= 10
    else:
        checks.append({"name": "段落结构", "status": "error", "message": "缺少段落划分，建议按工作经历/技能等分段", "score": 0})
        total_score -= 20
    
    # 5. 日期格式统一性检测
    date_patterns = [
        r'\d{4}年\d{1,2}月',           # 2023年5月
        r'\d{4}\.\d{1,2}',             # 2023.05
        r'\d{4}/\d{1,2}',              # 2023/05
        r'\d{1,2}/\d{4}',              # 05/2023
        r'(至今|present|now)',         # 至今
    ]
    dates_found = []
    for pattern in date_patterns:
        matches = re.findall(pattern, resume, re.IGNORECASE)
        dates_found.extend(matches)
    
    if len(dates_found) >= 2:
        # 检查格式是否一致（简化检测）
        has_chinese = any('年' in d or '月' in d or '至今' in d for d in dates_found)
        has_dot = any('.' in d for d in dates_found)
        if not (has_chinese and has_dot):  # 格式统一
            checks.append({"name": "日期格式", "status": "pass", "message": f"检测到{len(dates_found)}处日期，格式统一", "score": 100})
        else:
            checks.append({"name": "日期格式", "status": "warning", "message": "日期格式不统一，建议统一使用一种格式", "score": 70})
            total_score -= 5
    elif len(dates_found) == 1:
        checks.append({"name": "日期格式", "status": "warning", "message": "仅检测到1处日期，建议补充工作时间", "score": 60})
        total_score -= 10
    else:
        checks.append({"name": "日期格式", "status": "error", "message": "缺少日期信息，建议添加工作时间", "score": 0})
        total_score -= 15
    
    # 6. 特殊符号检测（ATS不友好）
    special_chars = r'[◆●■□▲►★☆※◎○◇△▽⊙⊕]'
    has_special = re.search(special_chars, resume)
    if has_special:
        special_found = re.findall(special_chars, resume)
        checks.append({"name": "特殊符号", "status": "warning", "message": f"包含{len(special_found)}个特殊符号（ATS可能无法识别）", "score": 80})
        total_score -= 5
    else:
        checks.append({"name": "特殊符号", "status": "pass", "message": "无特殊符号，ATS友好", "score": 100})
    
    # 7. 量化指标检测
    metrics_pattern = r'(提升|增长|增加|节省|减少|完成|达成)[^\d]*(\d+[%万千万亿]|[\d.]+[%万千万亿])'
    has_metrics = re.search(metrics_pattern, resume)
    if has_metrics:
        metrics_count = len(re.findall(metrics_pattern, resume))
        checks.append({"name": "量化成果", "status": "pass", "message": f"检测到{metrics_count}处量化成果数据", "score": 100})
    else:
        # 简化检测：只看数字+单位
        simple_metrics = re.findall(r'\d+[%万千万亿]', resume)
        if len(simple_metrics) >= 2:
            checks.append({"name": "量化成果", "status": "pass", "message": f"检测到{len(simple_metrics)}处数据指标", "score": 90})
        else:
            checks.append({"name": "量化成果", "status": "warning", "message": "缺少量化成果，建议添加具体数据（如：销售额增长30%）", "score": 50})
            total_score -= 10
    
    # 8. 行动动词检测
    action_verbs = ['负责', '完成', '实现', '优化', '开发', '管理', '协调', '策划', '执行', '分析', '提升', '建立', '主导', '推动', '达成', '获得', '荣获', '创造', '改善', '节约']
    verbs_found = [v for v in action_verbs if v in resume]
    if len(verbs_found) >= 5:
        checks.append({"name": "行动动词", "status": "pass", "message": f"使用{len(verbs_found)}个行动动词，表达积极", "score": 100})
    elif len(verbs_found) >= 3:
        checks.append({"name": "行动动词", "status": "warning", "message": f"仅{len(verbs_found)}个行动动词，建议增加", "score": 70})
        total_score -= 5
    else:
        checks.append({"name": "行动动词", "status": "error", "message": "缺少行动动词，建议使用如：负责、完成、优化等", "score": 0})
        total_score -= 10
    
    # 9. 技能关键词检测（常见技能词）
    skill_keywords = ['熟练', '精通', '掌握', '熟悉', '擅长', '具备', '了解', '运用']
    skills_found = [s for s in skill_keywords if s in resume]
    if len(skills_found) >= 3:
        checks.append({"name": "技能描述", "status": "pass", "message": f"包含{len(skills_found)}处技能描述词", "score": 100})
    elif len(skills_found) >= 1:
        checks.append({"name": "技能描述", "status": "warning", "message": f"仅{len(skills_found)}处技能描述，建议补充", "score": 70})
        total_score -= 5
    else:
        checks.append({"name": "技能描述", "status": "error", "message": "缺少技能熟练度描述词", "score": 0})
        total_score -= 10
    
    # 10. 工作经历检测（关键词）
    work_keywords = ['工作', '任职', '就职', '就职于', '岗位', '职位', '公司', '企业']
    work_found = [w for w in work_keywords if w in resume]
    if len(work_found) >= 2:
        checks.append({"name": "工作经历", "status": "pass", "message": f"检测到工作经历相关内容", "score": 100})
    elif len(work_found) >= 1:
        checks.append({"name": "工作经历", "status": "warning", "message": "工作经历描述较少，建议补充", "score": 70})
        total_score -= 5
    else:
        checks.append({"name": "工作经历", "status": "error", "message": "缺少工作经历描述", "score": 0})
        total_score -= 15
    
    # 11. 教育经历检测
    edu_keywords = ['大学', '学院', '学校', '学历', '专业', '本科', '硕士', '博士', '大专', '毕业']
    edu_found = [e for e in edu_keywords if e in resume]
    if len(edu_found) >= 2:
        checks.append({"name": "教育经历", "status": "pass", "message": f"检测到教育经历相关内容", "score": 100})
    elif len(edu_found) >= 1:
        checks.append({"name": "教育经历", "status": "warning", "message": "教育经历描述较少，建议补充", "score": 70})
        total_score -= 5
    else:
        checks.append({"name": "教育经历", "status": "error", "message": "缺少教育经历描述", "score": 0})
        total_score -= 10
    
    # 12. 简历标题/职位意向检测
    title_keywords = ['求职', '应聘', '意向', '目标职位', '期望', '应聘职位']
    has_title = any(t in resume[:200] for t in title_keywords)  # 检查前200字
    if has_title:
        checks.append({"name": "求职意向", "status": "pass", "message": "检测到求职意向/目标职位", "score": 100})
    else:
        checks.append({"name": "求职意向", "status": "warning", "message": "建议在开头添加求职意向", "score": 60})
        total_score -= 5
    
    # 计算最终分数（保底0分）
    final_score = max(0, min(100, total_score))
    
    return {
        "score": final_score,
        "checks": checks,
        "total_checks": 12,
        "passed": len([c for c in checks if c["status"] == "pass"]),
        "warnings": len([c for c in checks if c["status"] == "warning"]),
        "errors": len([c for c in checks if c["status"] == "error"])
    }

def generate_suggestions(resume: str, missing_keywords: list, industry: str):
    """生成优化建议"""
    suggestions = []
    industry_data = KEYWORDS_DB.get(industry, {})
    
    for kw in missing_keywords[:5]:
        suggestions.append({
            "type": "add_keyword",
            "keyword": kw["keyword"],
            "suggestion": f"建议添加「{kw['keyword']}」相关经验描述"
        })
    
    if not re.search(r'\d+%', resume) and not re.search(r'\d+万', resume):
        metrics = industry_data.get("metrics_examples", [])
        if metrics:
            suggestions.append({
                "type": "add_metric",
                "suggestion": f"建议添加量化数据，如：{metrics[0]}"
            })
    
    return suggestions

def generate_example_sentences(keyword: str, industry: str, context: str = "") -> list:
    """根据关键词和行业生成示例句子"""
    industry_name = KEYWORDS_DB.get(industry, {}).get("name", industry)
    
    # 行业特点模板库
    industry_templates = {
        "sales": {
            "prefixes": ["负责", "主导", "推动", "实现", "达成"],
            "suffixes": ["提升销售额XX%", "拓展XX个新客户", "建立XX渠道合作关系", "完成年度销售目标XX%"],
            "examples": [
                f"通过{keyword}策略，成功拓展了30+新客户，销售额同比增长25%",
                f"主导{keyword}项目，实现季度销售目标超额完成15%",
                f"运用{keyword}技巧，建立了50+渠道合作关系，市场覆盖率提升20%",
                f"负责{keyword}工作，年度业绩达成率120%，获得销售冠军称号"
            ]
        },
        "finance": {
            "prefixes": ["管理", "审核", "分析", "优化", "控制"],
            "suffixes": ["节省成本XX万元", "提高资金周转率XX%", "降低财务风险XX%", "优化财务流程"],
            "examples": [
                f"通过{keyword}分析，为公司节省成本15万元/年",
                f"负责{keyword}工作，财务报表准确率100%，零差错",
                f"运用{keyword}方法，优化资金配置，周转率提升30%",
                f"主导{keyword}项目，建立完善的内控体系，降低风险20%"
            ]
        },
        "administrative": {
            "prefixes": ["组织", "协调", "管理", "优化", "支持"],
            "suffixes": ["提升工作效率XX%", "服务XX人次", "处理XX事项", "优化流程"],
            "examples": [
                f"负责{keyword}工作，部门运转效率提升20%",
                f"通过{keyword}协调，成功组织XX场大型会议/活动",
                f"运用{keyword}技能，处理日常事务XX件，零延误",
                f"主导{keyword}优化，建立标准化流程，响应速度提升50%"
            ]
        },
        "customer_service": {
            "prefixes": ["服务", "解决", "处理", "维护", "提升"],
            "suffixes": ["满意度XX%", "响应时间XX分钟", "处理XX客户问题", "挽回XX客户"],
            "examples": [
                f"通过{keyword}服务，客户满意度达到95%",
                f"负责{keyword}工作，日均处理50+客户问题，解决率98%",
                f"运用{keyword}技巧，成功挽回流失客户30+",
                f"主导{keyword}改进，响应时间缩短至5分钟，投诉率下降40%"
            ]
        },
        "education": {
            "prefixes": ["教授", "指导", "培训", "开发", "设计"],
            "suffixes": ["学员XX人", "通过率XX%", "满意度XX%", "课程XX门"],
            "examples": [
                f"负责{keyword}教学，学员通过率95%，满意度4.8分",
                f"通过{keyword}方法，培训学员200+，就业率80%",
                f"运用{keyword}设计，开发精品课程10门，获好评",
                f"主导{keyword}项目，教学质量评估优秀，获教学奖"
            ]
        }
    }
    
    # 默认模板（通用）
    default_templates = {
        "prefixes": ["负责", "主导", "参与", "完成", "实现"],
        "suffixes": ["提升效率XX%", "完成XX项目", "达成XX目标", "获得XX成果"],
        "examples": [
            f"负责{keyword}工作，成功达成项目目标，效率提升20%",
            f"通过{keyword}实践，积累了丰富经验，取得显著成果",
            f"运用{keyword}技能，解决了XX难题，获得团队认可",
            f"主导{keyword}项目，按时完成交付，质量达标"
        ]
    }
    
    # 获取行业模板或使用默认
    templates = industry_templates.get(industry, default_templates)
    
    # 生成示例句子
    examples = templates.get("examples", default_templates["examples"])
    
    # 如果有上下文（如工作经历段落），可以结合生成更精准的句子
    if context:
        # 根据上下文判断句子类型
        if "负责" in context or "主管" in context:
            examples = [f"负责{keyword}工作，{templates['suffixes'][0]}" for kw_suffix in templates['suffixes'][:2]]
        elif "参与" in context or "协助" in context:
            examples = [f"参与{keyword}项目，协助完成目标" for _ in range(2)]
    
    # 返回多个示例供用户选择
    return examples[:4]  # 最多返回4个示例

def get_recommended_weights(industry: str) -> dict:
    """O5-2新增：根据行业推荐最佳权重配置"""
    # 行业权重推荐规则
    weight_rules = {
        # 销售类：关键词权重高（客户匹配更重要）
        "sales": {"keyword_weight": 70, "ats_weight": 30, "reason": "销售岗位更看重关键词匹配和客户相关经验"},
        "new_media": {"keyword_weight": 70, "ats_weight": 30, "reason": "新媒体岗位更看重内容关键词匹配"},
        "live_streaming": {"keyword_weight": 70, "ats_weight": 30, "reason": "直播岗位更看重关键词匹配"},
        "medical_sales": {"keyword_weight": 70, "ats_weight": 30, "reason": "医疗销售更看重专业关键词"},
        "real_estate_sales": {"keyword_weight": 70, "ats_weight": 30, "reason": "房产销售更看重关键词匹配"},
        "cross_border_ecommerce": {"keyword_weight": 65, "ats_weight": 35, "reason": "跨境电商关键词和格式都很重要"},
        
        # 技术类：ATS权重较高（格式和技术细节更重要）
        "it_developer": {"keyword_weight": 50, "ats_weight": 50, "reason": "技术岗位关键词和ATS格式同等重要"},
        "it_product": {"keyword_weight": 55, "ats_weight": 45, "reason": "产品岗位关键词稍重要"},
        "it_operation": {"keyword_weight": 50, "ats_weight": 50, "reason": "运维岗位关键词和格式同等重要"},
        
        # 财务类：均衡权重（精确性重要）
        "finance": {"keyword_weight": 55, "ats_weight": 45, "reason": "财务岗位关键词和格式都需要精确"},
        
        # 行政类：关键词权重稍高
        "administrative": {"keyword_weight": 60, "ats_weight": 40, "reason": "行政岗位关键词匹配较重要"},
        
        # 教育类：关键词权重高
        "education": {"keyword_weight": 65, "ats_weight": 35, "reason": "教育岗位关键词匹配较重要"},
        
        # 设计类：均衡权重
        "design": {"keyword_weight": 55, "ats_weight": 45, "reason": "设计岗位关键词和格式同等重要"},
        
        # 默认配置
        "default": {"keyword_weight": 60, "ats_weight": 40, "reason": "默认权重配置"}
    }
    
    # 获取推荐权重
    if industry in weight_rules:
        return weight_rules[industry]
    else:
        return weight_rules["default"]

def analyze_resume_structure(resume: str) -> dict:
    """分析简历结构，识别各段落区域"""
    # 定义常见简历段落标记
    section_markers = {
        "个人信息": ["个人信息", "联系方式", "电话", "邮箱", "姓名"],
        "求职意向": ["求职意向", "目标职位", "期望职位", "应聘职位"],
        "工作经历": ["工作经历", "工作经验", "任职经历", "工作履历", "就业经历"],
        "技能特长": ["技能", "专业技能", "特长", "技能特长", "能力", "核心技能"],
        "教育经历": ["教育经历", "教育背景", "学历", "毕业院校"],
        "项目经历": ["项目经历", "项目经验", "主要项目", "项目"],
        "自我评价": ["自我评价", "个人评价", "个人简介", "简介"]
    }
    
    sections = {}
    section_positions = []
    
    # 按段落分割简历
    paragraphs = resume.split('\n\n')
    
    for i, para in enumerate(paragraphs):
        para_clean = para.strip()
        if not para_clean:
            continue
        
        # 检测段落类型
        detected_type = None
        for section_type, markers in section_markers.items():
            for marker in markers:
                if marker in para_clean[:50]:  # 检查段落开头
                    detected_type = section_type
                    break
            if detected_type:
                break
        
        # 如果没检测到类型，根据内容特征判断
        if not detected_type:
            if re.search(r'\d{4}年|\d{4}\.\d+|至今', para_clean):
                if re.search(r'公司|企业|任职|就职', para_clean):
                    detected_type = "工作经历"
                elif re.search(r'大学|学院|学校|专业', para_clean):
                    detected_type = "教育经历"
            elif re.search(r'熟练|精通|掌握|熟悉|擅长', para_clean):
                detected_type = "技能特长"
        
        if detected_type:
            sections[detected_type] = {
                "index": i,
                "content": para_clean,
                "start_pos": resume.find(para_clean),
                "end_pos": resume.find(para_clean) + len(para_clean),
                "length": len(para_clean)
            }
            section_positions.append({
                "type": detected_type,
                "start": resume.find(para_clean),
                "end": resume.find(para_clean) + len(para_clean)
            })
    
    return {
        "sections": sections,
        "section_positions": section_positions,
        "total_paragraphs": len(paragraphs),
        "detected_types": list(sections.keys())
    }

def suggest_keyword_positions(resume: str, missing_keywords: list, industry: str) -> dict:
    """为缺失关键词推荐最佳插入位置"""
    structure = analyze_resume_structure(resume)
    
    # 关键词最佳位置规则（根据关键词类型）
    keyword_position_rules = {
        "技能类": ["技能特长", "工作经历", "自我评价"],
        "成果类": ["工作经历", "项目经历", "自我评价"],
        "经历类": ["工作经历", "项目经历"],
        "能力类": ["技能特长", "自我评价", "工作经历"],
        "通用类": ["工作经历", "技能特长", "自我评价", "求职意向"]
    }
    
    # 行业关键词类型映射（简化）
    industry_keyword_types = {
        "sales": {
            "客户开发": "技能类", "销售目标": "成果类", "业绩增长": "成果类",
            "渠道拓展": "技能类", "客户维护": "技能类"
        },
        "finance": {
            "成本控制": "技能类", "财务分析": "技能类", "预算管理": "技能类",
            "风险控制": "技能类", "报表编制": "技能类"
        },
        "administrative": {
            "流程优化": "成果类", "会议管理": "技能类", "文档管理": "技能类",
            "协调沟通": "能力类", "行政支持": "技能类"
        }
    }
    
    position_suggestions = []
    
    for kw in missing_keywords[:8]:  # 最多处理8个关键词
        kw_text = kw["keyword"]
        
        # 获取关键词类型
        kw_type = "通用类"
        if industry in industry_keyword_types:
            kw_type = industry_keyword_types[industry].get(kw_text, "通用类")
        
        # 获取推荐位置列表
        recommended_sections = keyword_position_rules.get(kw_type, keyword_position_rules["通用类"])
        
        # 检查简历中是否有这些段落
        available_sections = []
        for section_name in recommended_sections:
            if section_name in structure["sections"]:
                available_sections.append(section_name)
        
        # 如果没有匹配的段落，推荐工作经历或技能
        if not available_sections:
            if "工作经历" in structure["sections"]:
                available_sections = ["工作经历"]
            elif "技能特长" in structure["sections"]:
                available_sections = ["技能特长"]
            else:
                available_sections = ["简历末尾"]
        
        # 生成位置建议
        primary_section = available_sections[0]
        section_info = structure["sections"].get(primary_section, {})
        
        position_suggestions.append({
            "keyword": kw_text,
            "recommended_section": primary_section,
            "alternative_sections": available_sections[1:3] if len(available_sections) > 1 else [],
            "section_content_preview": section_info.get("content", "")[:100] if section_info else "",
            "insert_position": section_info.get("end_pos", 0) if section_info else len(resume),
            "suggestion_text": f"建议将「{kw_text}」添加到{primary_section}段落"
        })
    
    return {
        "suggestions": position_suggestions,
        "resume_structure": structure["detected_types"],
        "total_sections": len(structure["sections"])
    }

def get_keyword_synonyms(keyword: str, industry: str = "") -> list:
    """获取关键词同义词推荐"""
    # 同义词库（按行业分类）
    synonym_db = {
        "sales": {
            "客户开发": ["客户拓展", "客户获取", "新客户开发", "业务拓展"],
            "销售目标": ["销售业绩", "销售任务", "业绩指标", "销售配额"],
            "业绩增长": ["业绩提升", "销售增长", "业绩突破", "业务增长"],
            "渠道拓展": ["渠道开发", "渠道建设", "渠道布局", "销售渠道"],
            "客户维护": ["客户关系维护", "客户服务", "客户管理", "客户跟进"],
            "团队管理": ["团队建设", "团队领导", "人员管理", "团队培养"]
        },
        "finance": {
            "成本控制": ["成本管理", "成本优化", "成本降低", "费用控制"],
            "财务分析": ["财务报表分析", "数据分析", "财务诊断", "经营分析"],
            "预算管理": ["预算编制", "预算控制", "预算执行", "财务预算"],
            "风险控制": ["风险管理", "风险防范", "内控管理", "风险防控"],
            "报表编制": ["报表制作", "财务报告", "报表分析", "账务处理"]
        },
        "administrative": {
            "流程优化": ["流程改进", "流程再造", "流程梳理", "流程完善"],
            "会议管理": ["会议组织", "会议安排", "会议协调", "会务管理"],
            "文档管理": ["档案管理", "文件管理", "资料管理", "文档归档"],
            "协调沟通": ["沟通协调", "跨部门协作", "组织协调", "沟通联络"]
        },
        "tech": {
            "项目管理": ["项目推进", "项目执行", "项目实施", "项目统筹"],
            "团队协作": ["团队合作", "协同工作", "跨部门协作", "团队配合"],
            "技术优化": ["性能优化", "技术改进", "代码优化", "系统优化"],
            "需求分析": ["需求调研", "需求梳理", "业务分析", "需求评估"]
        },
        "marketing": {
            "市场推广": ["市场拓展", "品牌推广", "市场宣传", "营销推广"],
            "活动策划": ["活动组织", "营销策划", "活动执行", "方案策划"],
            "内容运营": ["内容创作", "内容管理", "内容营销", "内容制作"],
            "数据分析": ["数据挖掘", "数据统计", "数据报告", "数据分析"]
        }
    }
    
    # 通用同义词
    common_synonyms = {
        "负责": ["主管", "承担", "主导", "统筹"],
        "参与": ["协助", "配合", "参与", "加入"],
        "完成": ["达成", "实现", "完成", "取得"],
        "提升": ["提高", "优化", "增强", "改进"],
        "管理": ["管理", "管控", "治理", "监管"],
        "优化": ["改进", "完善", "优化", "提升"],
        "协调": ["协调", "联络", "沟通", "对接"]
    }
    
    synonyms = []
    
    # 1. 行业特定同义词
    if industry in synonym_db:
        industry_synonyms = synonym_db[industry].get(keyword, [])
        synonyms.extend(industry_synonyms[:3])  # 最多3个行业同义词
    
    # 2. 通用同义词
    for key, values in common_synonyms.items():
        if key in keyword and keyword != key:
            # 为关键词生成变体
            for v in values[:2]:
                variant = keyword.replace(key, v)
                if variant != keyword and variant not in synonyms:
                    synonyms.append(variant)
    
    # 3. 如果没有找到同义词，生成一些通用的变体
    if not synonyms:
        # 添加"能力"、"经验"等后缀变体
        if not keyword.endswith("能力"):
            synonyms.append(f"{keyword}能力")
        if not keyword.endswith("经验"):
            synonyms.append(f"{keyword}经验")
    
    return synonyms[:6]  # 返回最多6个同义词

def analyze_keyword_density(resume: str, keywords_found: list, all_keywords: list) -> dict:
    """关键词密度分析"""
    # 计算总字数
    total_chars = len(resume)
    if total_chars == 0:
        return {"density": 0, "distribution": "empty", "warnings": []}
    
    # 计算关键词总字数
    keyword_chars = 0
    keyword_positions = []  # 记录关键词位置
    
    for kw in keywords_found:
        kw_text = kw["keyword"]
        kw_len = len(kw_text)
        keyword_chars += kw_len
        
        # 查找关键词位置
        pos = resume.find(kw_text)
        if pos != -1:
            keyword_positions.append({
                "keyword": kw_text,
                "position": pos,
                "section": pos // (total_chars // 5) if total_chars > 0 else 0  # 分5段
            })
    
    # 计算密度百分比
    density = round(keyword_chars / total_chars * 100, 2) if total_chars > 0 else 0
    
    # 判断密度是否合理（推荐2-5%）
    density_status = "normal"
    density_warnings = []
    
    if density < 2:
        density_status = "low"
        density_warnings.append("关键词密度过低（<2%），建议增加行业关键词")
    elif density > 8:
        density_status = "high"
        density_warnings.append("关键词密度过高（>8%），可能存在堆砌风险")
    elif density > 5:
        density_status = "slightly_high"
        density_warnings.append("关键词密度偏高，注意避免堆砌")
    
    # 分析关键词分布（是否分散在各段落）
    if keyword_positions:
        # 将简历分成5段
        sections = [0, 0, 0, 0, 0]
        for kp in keyword_positions:
            section_idx = min(kp["section"], 4)
            sections[section_idx] += 1
        
        # 计算分布均匀度
        total_kw = len(keyword_positions)
        distribution_status = "balanced"
        
        if total_kw > 0:
            max_section = max(sections)
            min_section = min(sections)
            
            # 如果某个段落关键词过多
            if max_section > total_kw * 0.5:
                distribution_status = "concentrated"
                density_warnings.append(f"关键词集中在第{sections.index(max_section)+1}部分，建议分散到各段落")
            # 如果某些段落没有关键词
            elif min_section == 0 and total_kw >= 3:
                empty_sections = [i for i, s in enumerate(sections) if s == 0]
                distribution_status = "unbalanced"
                density_warnings.append(f"第{[i+1 for i in empty_sections]}部分缺少关键词，建议补充")
    else:
        distribution_status = "no_keywords"
        density_warnings.append("未检测到行业关键词")
    
    return {
        "density": density,
        "density_status": density_status,
        "distribution": distribution_status,
        "keyword_positions": keyword_positions[:10],  # 只返回前10个
        "warnings": density_warnings,
        "total_keywords": len(keywords_found),
        "total_chars": total_chars,
        "recommendation": "建议关键词密度保持在2-5%，均匀分布在各段落"
    }

def analyze_quantified_metrics(resume: str) -> dict:
    """量化指标检测与分析"""
    # 量化指标模式（数字+单位）
    patterns = {
        "百分比": r'\d+(?:\.\d+)?%',
        "金额": r'\d+(?:\.\d+)?(?:万|千万|亿|元|块)',
        "人数": r'\d+(?:\.\d+)?(?:人|名|位|个)',
        "时间": r'\d+(?:\.\d+)?(?:年|月|天|小时|分钟|次)',
        "数量": r'\d+(?:\.\d+)?(?:个|件|条|篇|份|户|客户|订单)'
    }
    
    # 更复杂的量化模式（动词+数字+单位）
    action_patterns = [
        r'(提升|增长|增加|提高|上升)\s*[^\d]*\d+(?:\.\d+)?%',
        r'(节省|减少|降低|下降|节约)\s*[^\d]*\d+(?:\.\d+)?(?:万|千万|亿|元|%)',
        r'(完成|达成|实现|达到)\s*[^\d]*\d+(?:\.\d+)?(?:万|千万|亿|%|个|件)',
        r'(管理|带领|负责|领导)\s*[^\d]*\d+(?:\.\d+)?(?:人|名|位)',
        r'(处理|服务|接待|响应)\s*[^\d]*\d+(?:\.\d+)?(?:个|件|户|客户|订单)'
    ]
    
    metrics_found = []
    total_metrics = 0
    
    # 按类型检测
    for metric_type, pattern in patterns.items():
        matches = re.findall(pattern, resume)
        if matches:
            for match in matches[:5]:  # 每类最多5个
                metrics_found.append({
                    "type": metric_type,
                    "value": match,
                    "context": ""  # 可以后续提取上下文
                })
            total_metrics += len(matches)
    
    # 检测动词+数值模式（更专业）
    action_metrics = []
    for pattern in action_patterns:
        matches = re.findall(pattern, resume)
        for match in matches:
            action_metrics.append(match)
    
    # 评分和建议
    score = 0
    warnings = []
    suggestions = []
    
    if total_metrics >= 5:
        score = 100
    elif total_metrics >= 3:
        score = 80
        warnings.append("量化数据较少，建议增加更多具体成果数据")
    elif total_metrics >= 1:
        score = 50
        warnings.append("仅检测到少量量化数据，建议补充更多成果指标")
    else:
        score = 0
        warnings.append("缺少量化成果数据，强烈建议添加具体数字指标")
    
    # 检测工作经历段落是否缺少量化
    work_sections = re.findall(r'(?:工作经历|工作经验|任职|就职)[^\n]*\n(?:[^\n]+\n){1,5}', resume)
    for i, section in enumerate(work_sections):
        if not re.search(r'\d+', section):
            warnings.append(f"工作经历第{i+1}段缺少量化数据，建议添加具体成果")
            suggestions.append({
                "type": "add_metric",
                "section": f"工作经历第{i+1}段",
                "example": "例如：销售额增长30%、节省成本10万元、管理团队15人"
            })
    
    # 量化示例模板
    metric_templates = [
        "销售额增长XX%",
        "节省成本XX万元",
        "管理团队XX人",
        "处理XX个客户/订单",
        "提升效率XX%",
        "完成XX个项目"
    ]
    
    return {
        "score": score,
        "total_metrics": total_metrics,
        "metrics_found": metrics_found[:15],  # 最多返回15个
        "action_metrics": action_metrics[:5],  # 动词+数值模式
        "warnings": warnings,
        "suggestions": suggestions,
        "templates": metric_templates,
        "recommendation": "简历中至少应包含3-5个量化成果数据，展示具体业绩"
    }

# ========== 其他API ==========

@app.get("/")
async def root():
    return {
        "name": "ResumeAI API",
        "version": "2.0.0",
        "industries": list(KEYWORDS_DB.keys()),
        "auth_enabled": True
    }

@app.get("/api/v1/industries")
async def get_industries():
    industries = []
    for key, data in KEYWORDS_DB.items():
        industries.append({
            "id": key,
            "name": data.get("name", key),
            "name_en": data.get("name_en", key),
            "has_sub_industries": "sub_industries" in data  # O4-1：标识是否有子行业
        })
    return {"success": True, "data": industries}

# O4-1新增：获取子行业列表
@app.get("/api/v1/sub-industries/{industry_id}")
async def get_sub_industries(industry_id: str):
    """获取指定行业的子行业列表"""
    industry_data = KEYWORDS_DB.get(industry_id)
    if not industry_data:
        return {"success": False, "error": "行业不存在"}
    
    sub_industries = industry_data.get("sub_industries", [])
    if not sub_industries:
        return {"success": True, "data": [], "message": "该行业暂无子行业分类"}
    
    result = []
    for sub in sub_industries:
        result.append({
            "id": sub.get("id"),
            "name": sub.get("name"),
            "keywords_count": len(sub.get("keywords_boost", []))
        })
    
    return {"success": True, "data": result}

# O5-2新增：获取行业推荐权重
@app.get("/api/v1/recommended-weights/{industry}")
async def get_weights_recommendation(industry: str):
    """获取指定行业的推荐评分权重配置"""
    recommendation = get_recommended_weights(industry)
    return {
        "success": True,
        "industry": industry,
        "recommended_keyword_weight": recommendation["keyword_weight"],
        "recommended_ats_weight": recommendation["ats_weight"],
        "reason": recommendation["reason"]
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "keywords_loaded": len(KEYWORDS_DB),
        "auth_enabled": True,
        "data_dir": DATA_DIR,
        "db_path": USER_DB_PATH,
        "version": "2026-05-16-v7"  # 版本标记（添加admin/trends API）
    }

@app.get("/debug/db")
async def debug_db():
    """调试数据库状态"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        # 检查表
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        # 检查users表结构
        cursor.execute("PRAGMA table_info(users)")
        users_columns = [row[1] for row in cursor.fetchall()]
        # 检查用户数量
        cursor.execute("SELECT COUNT(*) FROM users")
        user_count = cursor.fetchone()[0]
        conn.close()
        return {
            "tables": tables,
            "users_columns": users_columns,
            "user_count": user_count,
            "db_path": USER_DB_PATH,
            "data_dir": DATA_DIR
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/debug/init-tables")
async def init_tables():
    """初始化/修复数据库表结构"""
    try:
        conn = sqlite3.connect(USER_DB_PATH)
        cursor = conn.cursor()
        
        # 检查users表结构
        cursor.execute("PRAGMA table_info(users)")
        columns = [row[1] for row in cursor.fetchall()]
        
        # 添加缺失的字段
        missing_columns = []
        
        if 'is_admin' not in columns:
            cursor.execute("ALTER TABLE users ADD COLUMN is_admin INTEGER DEFAULT 0")
            missing_columns.append('is_admin')
        
        if 'is_paid' not in columns:
            cursor.execute("ALTER TABLE users ADD COLUMN is_paid INTEGER DEFAULT 0")
            missing_columns.append('is_paid')
        
        # 创建verification_codes表
        cursor.execute('''CREATE TABLE IF NOT EXISTS verification_codes
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT,
            code TEXT,
            expires_at TEXT,
            used INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP)''')
        
        conn.commit()
        
        # 检查最终表结构
        cursor.execute("PRAGMA table_info(users)")
        users_columns = [row[1] for row in cursor.fetchall()]
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        conn.close()
        
        return {
            "success": True, 
            "tables": tables,
            "users_columns": users_columns,
            "added_columns": missing_columns,
            "message": "数据库表结构已修复"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/debug/create-admin")
async def create_admin(
    email: str = Query(default="zhwffy@hotmail.com"),
    name: str = Query(default="Admin"),
    password: str = Query(default="Hiller")
):
    """创建管理员账号"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        
        # 检查用户是否存在
        cursor.execute("SELECT id FROM users WHERE email = ?", (email,))
        if cursor.fetchone():
            # 更新为管理员
            cursor.execute("UPDATE users SET is_admin = 1 WHERE email = ?", (email,))
            conn.commit()
            conn.close()
            return {"success": True, "message": "已将用户设置为管理员", "email": email}
        
        # 创建新管理员
        password_hash = hash_password(password)
        cursor.execute('''INSERT INTO users 
            (email, name, password_hash, is_admin, is_paid, created_at)
            VALUES (?, ?, ?, 1, 1, ?)''', 
            (email, name, password_hash, datetime.now().isoformat()))
        conn.commit()
        user_id = cursor.lastrowid
        conn.close()
        
        # 生成token
        token = create_token(user_id, email)
        
        return {
            "success": True,
            "message": "管理员账号已创建",
            "user": {
                "id": user_id,
                "email": email,
                "name": name,
                "is_admin": True
            },
            "token": token
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/debug/test-email")
async def test_email(request: SendCodeRequest):
    """测试邮件发送"""
    try:
        code = generate_code(EMAIL_CONFIG["code_length"])
        result = send_verification_code(request.email, code)
        return {
            "success": result["success"],
            "message": result["message"],
            "test_mode": EMAIL_CONFIG["test_mode"],
            "smtp_server": EMAIL_CONFIG["smtp_server"],
            "smtp_user": EMAIL_CONFIG["smtp_user"],
            "error": result.get("error") if not result["success"] else None
        }
    except Exception as e:
        return {"success": False, "error": str(e), "error_type": type(e).__name__}

@app.get("/debug/users")
async def debug_users():
    """查看所有用户"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        cursor.execute("SELECT id, email, name, password_hash, created_at FROM users")
        users = []
        for row in cursor.fetchall():
            users.append({
                "id": row["id"],
                "email": row["email"],
                "name": row["name"],
                "password_hash": row["password_hash"],
                "created_at": row["created_at"]
            })
        conn.close()
        return {"users": users, "count": len(users)}
    except Exception as e:
        return {"error": str(e)}

@app.post("/debug/check-password")
async def debug_check_password(email: str, password: str):
    """检查密码匹配"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()
        conn.close()
        
        if not user:
            return {"error": "用户不存在", "email": email}
        
        stored_hash = user["password_hash"]
        input_hash = hash_password(password)
        
        return {
            "email": email,
            "stored_hash": stored_hash,
            "input_hash": input_hash,
            "match": stored_hash == input_hash,
            "password": password
        }
    except Exception as e:
        return {"error": str(e)}

# ========== 简历模板API ==========

RESUME_TEMPLATES_FILE = os.path.join(os.path.dirname(__file__), "resume_templates.json")

def load_templates():
    """加载简历模板"""
    try:
        with open(RESUME_TEMPLATES_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return {"templates": {}, "default_template": "classic"}

@app.get("/api/v1/templates")
async def get_templates():
    """获取所有简历模板"""
    templates_data = load_templates()
    return {
        "success": True,
        "templates": templates_data["templates"],
        "default": templates_data.get("default_template", "classic")
    }

@app.get("/api/v1/templates/{template_id}")
async def get_template(template_id: str):
    """获取单个模板详情"""
    templates_data = load_templates()
    template = templates_data["templates"].get(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    return {"success": True, "template": template}

@app.get("/api/v1/templates/recommend/{industry}")
async def recommend_template(industry: str):
    """根据行业推荐模板"""
    templates_data = load_templates()
    recommended = []
    for tid, template in templates_data["templates"].items():
        if industry in template.get("suitable_for", []):
            recommended.append({"id": tid, "name": template["name"], "match_score": 100})
    
    if not recommended:
        # 使用默认模板
        default_tid = templates_data.get("default_template", "classic")
        default_template = templates_data["templates"].get(default_tid)
        if default_template:
            recommended.append({"id": default_tid, "name": default_template["name"], "match_score": 50})
    
    return {
        "success": True,
        "industry": industry,
        "recommended": recommended
    }

@app.get("/api/v1/templates/{template_id}/download")
async def download_template(template_id: str):
    """下载空白简历模板"""
    from fastapi.responses import FileResponse
    import os
    
    templates_dir = os.path.join(os.path.dirname(__file__), "templates")
    templates_file = os.path.join(templates_dir, "templates.json")
    
    # 读取模板配置
    with open(templates_file, "r", encoding="utf-8") as f:
        templates_data = json.load(f)
    
    template = templates_data["templates"].get(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = os.path.join(templates_dir, template["file"])
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    return FileResponse(
        path=file_path,
        filename=template["file"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ========== 用户模板存储API ==========

class SaveTemplateRequest(BaseModel):
    template_id: str
    custom_settings: Optional[str] = "{}"  # JSON格式的自定义设置
    is_default: Optional[bool] = False

@app.post("/api/v1/user/templates")
async def save_user_template(request: SaveTemplateRequest, user: dict = Depends(get_current_user)):
    """保存用户模板偏好"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    user_id = user["id"]  # 使用正确的字段名
    
    # 检查是否已存在
    cursor.execute("SELECT id FROM user_templates WHERE user_id = ? AND template_id = ?", 
                   (user_id, request.template_id))
    existing = cursor.fetchone()
    
    if existing:
        # 更新
        cursor.execute('''UPDATE user_templates 
            SET custom_settings = ?, is_default = ?, updated_at = ?
            WHERE id = ?''',
            (request.custom_settings, 1 if request.is_default else 0, 
             datetime.now().isoformat(), existing["id"]))
    else:
        # 新建
        cursor.execute('''INSERT INTO user_templates 
            (user_id, template_id, custom_settings, is_default, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)''',
            (user_id, request.template_id, request.custom_settings,
             1 if request.is_default else 0, datetime.now().isoformat(), datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "模板已保存"}

@app.get("/api/v1/user/templates")
async def get_user_templates(user: dict = Depends(get_current_user)):
    """获取用户保存的模板列表"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute('''SELECT id, template_id, custom_settings, is_default, created_at 
        FROM user_templates WHERE user_id = ?''', (user["id"],))
    
    templates = []
    for row in cursor.fetchall():
        templates.append({
            "id": row["id"],
            "template_id": row["template_id"],
            "custom_settings": row["custom_settings"],
            "is_default": row["is_default"],
            "created_at": row["created_at"]
        })
    
    conn.close()
    return {"success": True, "templates": templates}

@app.delete("/api/v1/user/templates/{template_id}")
async def delete_user_template(template_id: int, user: dict = Depends(get_current_user)):
    """删除用户模板"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("DELETE FROM user_templates WHERE id = ? AND user_id = ?", 
                   (template_id, user["id"]))
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "模板已删除"}

# ========== 调试API ==========

@app.get("/api/v1/debug/status")
async def debug_status():
    """调试状态检查"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as count FROM users")
    users_count = cursor.fetchone()["count"]
    cursor.execute("SELECT email FROM users LIMIT 10")
    users_emails = [row["email"] for row in cursor.fetchall()]
    conn.close()
    
    return {
        "success": True,
        "keywords_loaded": len(KEYWORDS_DB),
        "keywords_industries": list(KEYWORDS_DB.keys())[:10] if KEYWORDS_DB else [],
        "data_dir": DATA_DIR,
        "user_db_path": USER_DB_PATH,
        "has_sales": "sales" in KEYWORDS_DB,
        "jwt_secret": JWT_SECRET[:10] + "...",
        "users_count": users_count,
        "users_emails": users_emails
    }

@app.post("/api/v1/debug/analyze")
async def debug_analyze():
    """调试简历分析"""
    try:
        # 测试分析函数
        test_resume = "张三销售经理电话13800138000"
        result = analyze_resume(test_resume, "sales", "zh", None)
        return {"success": True, "result": result}
    except Exception as e:
        import traceback
        return {
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        }

# ========== 支付API (LemonSqueezy) ==========

# LemonSqueezy配置（占位，需用户创建产品后替换）
LEMONSQUEEZY_API_KEY = os.getenv("LEMONSQUEEZY_API_KEY", "")
LEMONSQUEEZY_STORE_ID = os.getenv("LEMONSQUEEZY_STORE_ID", "")
LEMONSQUEEZY_WEBHOOK_SECRET = os.getenv("LEMONSQUEEZY_WEBHOOK_SECRET", "")

class CreateCheckoutRequest(BaseModel):
    product_id: str
    variant_id: Optional[str] = None
    email: Optional[str] = None

@app.get("/api/v1/payments/status")
async def payment_status():
    """检查支付配置状态"""
    return {
        "success": True,
        "configured": bool(LEMONSQUEEZY_API_KEY and LEMONSQUEEZY_STORE_ID),
        "store_id": LEMONSQUEEZY_STORE_ID if LEMONSQUEEZY_API_KEY else None,
        "message": "支付功能配置完成" if LEMONSQUEEZY_API_KEY else "请配置LemonSqueezy API Key和Store ID"
    }

@app.post("/api/v1/payments/checkout")
async def create_checkout(request: CreateCheckoutRequest, user: dict = Depends(get_current_user)):
    """创建支付checkout链接"""
    if not LEMONSQUEEZY_API_KEY or not LEMONSQUEEZY_STORE_ID:
        return {"error": "支付功能未配置，请联系管理员"}
    
    try:
        # LemonSqueezy Checkout API调用
        import httpx
        
        payload = {
            "data": {
                "type": "checkouts",
                "attributes": {
                    "checkout_data": {
                        "email": request.email or user["email"],
                        "custom": {
                            "user_id": str(user["id"])
                        }
                    }
                },
                "relationships": {
                    "store": {"type": "stores", "id": LEMONSQUEEZY_STORE_ID},
                    "variant": {"type": "variants", "id": request.variant_id or request.product_id}
                }
            }
        }
        
        headers = {
            "Authorization": f"Bearer {LEMONSQUEEZY_API_KEY}",
            "Content-Type": "application/vnd.api+json"
        }
        
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                "https://api.lemonsqueezy.com/v1/checkouts",
                json=payload,
                headers=headers
            )
            
            if resp.status_code == 201:
                data = resp.json()
                checkout_url = data["data"]["attributes"]["url"]
                return {"success": True, "checkout_url": checkout_url}
            else:
                logger.error(f"LemonSqueezy checkout failed: {resp.text}")
                return {"error": "创建支付链接失败"}
    
    except Exception as e:
        logger.error(f"Payment error: {e}")
        return {"error": str(e)}

@app.post("/api/v1/payments/webhook")
async def payment_webhook(request: dict, x_signature: Optional[str] = Header(None)):
    """处理LemonSqueezy webhook回调"""
    # 验证签名
    if LEMONSQUEEZY_WEBHOOK_SECRET and x_signature:
        # TODO: 实现签名验证
        pass
    
    event_type = request.get("meta", {}).get("event_name", "")
    
    if event_type == "order_created":
        # 订单创建成功
        order_data = request["data"]["attributes"]
        custom_data = request["meta"]["custom_data"] or {}
        user_id = custom_data.get("user_id")
        
        if user_id:
            # 更新用户为付费用户
            conn = get_user_db()
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE users SET is_paid = 1 WHERE id = ?",
                (int(user_id),)
            )
            
            # 记录订单
            cursor.execute(
                "INSERT INTO orders (user_id, order_id, amount, status, created_at) VALUES (?, ?, ?, ?, ?)",
                (int(user_id), order_data["order_id"], order_data["total"], "completed", datetime.now().isoformat())
            )
            conn.commit()
            conn.close()
            
            logger.info(f"Payment successful for user {user_id}")
    
    return {"success": True}

@app.get("/api/v1/user/orders")
async def get_user_orders(user: dict = Depends(get_current_user)):
    """获取用户订单历史"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 确保orders表存在
    cursor.execute('''CREATE TABLE IF NOT EXISTS orders
        (id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        order_id TEXT,
        amount REAL,
        status TEXT,
        created_at TEXT)''')
    
    cursor.execute(
        "SELECT * FROM orders WHERE user_id = ? ORDER BY created_at DESC",
        (user["id"],)
    )
    
    orders = []
    for row in cursor.fetchall():
        orders.append({
            "id": row["id"],
            "order_id": row["order_id"],
            "amount": row["amount"],
            "status": row["status"],
            "created_at": row["created_at"]
        })
    
    conn.close()
    return {"success": True, "orders": orders}

# ========== 管理员API ==========

async def get_admin_user(authorization: Optional[str] = Header(None)):
    """验证管理员权限"""
    if not authorization:
        raise HTTPException(status_code=401, detail="未提供认证token")
    
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    payload = verify_token(token)
    
    if not payload:
        raise HTTPException(status_code=401, detail="token无效或已过期")
    
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (payload["user_id"],))
    user = cursor.fetchone()
    conn.close()
    
    if not user:
        raise HTTPException(status_code=401, detail="用户不存在")
    
    if not user["is_admin"]:
        raise HTTPException(status_code=403, detail="需要管理员权限")
    
    return dict(user)

@app.get("/api/v1/admin/stats")
async def admin_stats(admin: dict = Depends(get_admin_user)):
    """获取系统统计"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 用户统计
    cursor.execute("SELECT COUNT(*) as count FROM users")
    row = cursor.fetchone()
    total_users = row["count"] if row else 0
    
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE is_paid = 1")
    row = cursor.fetchone()
    paid_users = row["count"] if row else 0
    
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE DATE(created_at) = DATE('now')")
    row = cursor.fetchone()
    new_users_today = row["count"] if row else 0
    
    # 等级分布
    cursor.execute("SELECT user_level, COUNT(*) as count FROM users GROUP BY user_level")
    level_distribution = {}
    for row in cursor.fetchall():
        level_distribution[row["user_level"] or "free"] = row["count"]
    
    # VIP用户数
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE user_level = 'vip'")
    row = cursor.fetchone()
    vip_users = row["count"] if row else 0
    
    # 使用统计
    cursor.execute("SELECT COUNT(*) as count FROM usage")
    row = cursor.fetchone()
    total_usage = row["count"] if row else 0
    
    cursor.execute("SELECT COUNT(*) as count FROM usage WHERE DATE(created_at) = DATE('now')")
    row = cursor.fetchone()
    usage_today = row["count"] if row else 0
    
    # 简历统计
    cursor.execute("SELECT COUNT(*) as count FROM resumes")
    row = cursor.fetchone()
    total_resumes = row["count"] if row else 0
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "total_users": total_users,
            "paid_users": paid_users,
            "vip_users": vip_users,
            "new_users_today": new_users_today,
            "total_usage": total_usage,
            "today_usage": usage_today,
            "total_resumes": total_resumes,
            "level_distribution": level_distribution
        }
    }

# ========== O4-2新增：用户自定义关键词API ==========

@app.post("/api/v1/user-keywords")
async def add_user_keyword(request: AddKeywordRequest, user: dict = Depends(get_current_user)):
    """添加用户自定义关键词"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 检查是否已存在
    cursor.execute("""
        SELECT id FROM user_keywords WHERE user_id = ? AND keyword = ? AND industry = ?
    """, (user["id"], request.keyword, request.industry))
    
    if cursor.fetchone():
        return {"success": False, "error": "关键词已存在"}
    
    # 添加关键词
    cursor.execute("""
        INSERT INTO user_keywords (user_id, keyword, industry, weight, created_at)
        VALUES (?, ?, ?, ?, ?)
    """, (user["id"], request.keyword, request.industry, request.weight, datetime.now().isoformat()))
    conn.commit()
    
    return {"success": True, "message": "关键词添加成功", "keyword": request.keyword, "weight": request.weight}

@app.get("/api/v1/user-keywords/{industry}")
async def get_user_keywords(industry: str, user: dict = Depends(get_current_user)):
    """获取用户指定行业的自定义关键词"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT keyword, weight FROM user_keywords WHERE user_id = ? AND industry = ?
        ORDER BY weight DESC
    """, (user["id"], industry))
    
    keywords = [{"keyword": row[0], "weight": row[1]} for row in cursor.fetchall()]
    
    return {"success": True, "keywords": keywords}

@app.delete("/api/v1/user-keywords/{keyword_id}")
async def delete_user_keyword(keyword_id: int, user: dict = Depends(get_current_user)):
    """删除用户自定义关键词"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        DELETE FROM user_keywords WHERE id = ? AND user_id = ?
    """, (keyword_id, user["id"]))
    conn.commit()
    
    if cursor.rowcount > 0:
        return {"success": True, "message": "关键词删除成功"}
    else:
        return {"success": False, "error": "关键词不存在或无权限删除"}

@app.post("/api/v1/feedback")
async def submit_feedback(request: FeedbackRequest, user: Optional[dict] = Depends(get_current_user_optional)):
    """提交用户反馈（增强版）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    user_id = user["id"] if user else None
    email = request.email or (user["email"] if user else "anonymous")
    title = request.title or request.content[:50]  # 默认取内容前50字作为标题
    tags_str = json.dumps(request.tags) if request.tags else None
    
    cursor.execute("""
        INSERT INTO feedback (user_id, email, type, title, content, priority, tags, screenshot_url, status, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)
    """, (user_id, email, request.type, title, request.content, request.priority, tags_str, request.screenshot, datetime.now().isoformat()))
    conn.commit()
    feedback_id = cursor.lastrowid
    
    logger.info(f"新反馈提交: #{feedback_id} 类型:{request.type} 优先级:{request.priority}")
    
    return {
        "success": True, 
        "message": "反馈已提交，感谢您的建议！",
        "feedback_id": feedback_id
    }

@app.get("/api/v1/feedback/my")
async def get_my_feedback(user: dict = Depends(get_current_user)):
    """获取用户自己的反馈列表"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT id, type, title, content, priority, tags, status, admin_reply, created_at, reply_at
        FROM feedback WHERE user_id = ?
        ORDER BY created_at DESC
        LIMIT 50
    """, (user["id"],))
    
    feedbacks = []
    for row in cursor.fetchall():
        tags = json.loads(row["tags"]) if row["tags"] else []
        feedbacks.append({
            "id": row["id"],
            "type": row["type"],
            "title": row["title"],
            "content": row["content"][:100] + "..." if len(row["content"]) > 100 else row["content"],
            "priority": row["priority"],
            "tags": tags,
            "status": row["status"],
            "admin_reply": row["admin_reply"],
            "created_at": row["created_at"],
            "reply_at": row["reply_at"]
        })
    
    return {"success": True, "data": feedbacks}

@app.get("/api/v1/feedback/{feedback_id}")
async def get_feedback_detail(feedback_id: int, user: dict = Depends(get_current_user)):
    """获取反馈详情"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT id, type, title, content, priority, tags, screenshot_url, status, admin_reply, created_at, reply_at, rating
        FROM feedback WHERE id = ? AND user_id = ?
    """, (feedback_id, user["id"]))
    
    row = cursor.fetchone()
    if not row:
        return {"success": False, "message": "反馈不存在"}
    
    tags = json.loads(row["tags"]) if row["tags"] else []
    
    return {
        "success": True,
        "data": {
            "id": row["id"],
            "type": row["type"],
            "title": row["title"],
            "content": row["content"],
            "priority": row["priority"],
            "tags": tags,
            "screenshot_url": row["screenshot_url"],
            "status": row["status"],
            "admin_reply": row["admin_reply"],
            "created_at": row["created_at"],
            "reply_at": row["reply_at"],
            "rating": row["rating"]
        }
    }

@app.get("/api/v1/admin/feedback")
async def admin_get_feedback(admin: dict = Depends(get_admin_user)):
    """获取反馈列表（管理员，增强版）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT f.id, f.user_id, f.email, f.type, f.title, f.content, f.priority, f.tags, 
               f.screenshot_url, f.status, f.admin_reply, f.created_at, f.reply_at,
               u.name as user_name
        FROM feedback f
        LEFT JOIN users u ON f.user_id = u.id
        ORDER BY 
            CASE f.priority WHEN 'urgent' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END,
            f.created_at DESC
        LIMIT 100
    """)
    
    feedbacks = []
    for row in cursor.fetchall():
        tags = json.loads(row["tags"]) if row["tags"] else []
        feedbacks.append({
            "id": row["id"],
            "user_id": row["user_id"],
            "email": row["email"],
            "type": row["type"],
            "title": row["title"],
            "content": row["content"][:200] if row["content"] else "",
            "priority": row["priority"],
            "tags": tags,
            "screenshot_url": row["screenshot_url"],
            "status": row["status"],
            "admin_reply": row["admin_reply"],
            "created_at": row["created_at"],
            "reply_at": row["reply_at"],
            "user_name": row["user_name"]
        })
    
    return {"success": True, "data": feedbacks}

@app.put("/api/v1/admin/feedback/{feedback_id}")
async def admin_update_feedback(feedback_id: int, status: str = None, reply: str = None, admin: dict = Depends(get_admin_user)):
    """更新反馈（管理员，支持回复+通知）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取反馈信息
    cursor.execute("SELECT * FROM feedback WHERE id = ?", (feedback_id,))
    feedback = cursor.fetchone()
    if not feedback:
        conn.close()
        return {"success": False, "message": "反馈不存在"}
    
    old_status = feedback["status"]
    
    if status:
        cursor.execute("UPDATE feedback SET status = ? WHERE id = ?", (status, feedback_id))
        
        # 状态变化通知
        status_text = {"pending": "待处理", "in_progress": "处理中", "resolved": "已解决", "closed": "已关闭"}
        if feedback["user_id"] and status != old_status:
            create_notification(
                user_id=feedback["user_id"],
                title=f"反馈状态更新 #{feedback_id}",
                message=f"您的反馈「{feedback['title'] or '无标题'}」状态已更新为：{status_text.get(status, status)}",
                notif_type="feedback",
                related_id=feedback_id,
                related_type="feedback"
            )
    
    if reply:
        cursor.execute("""
            UPDATE feedback SET admin_reply = ?, reply_at = ?, status = 'in_progress'
            WHERE id = ?
        """, (reply, datetime.now().isoformat(), feedback_id))
        
        # 回复通知
        if feedback["user_id"]:
            create_notification(
                user_id=feedback["user_id"],
                title=f"您收到管理员回复 #{feedback_id}",
                message=f"管理员回复了您的反馈「{feedback['title'] or '无标题'}」",
                notif_type="feedback_reply",
                related_id=feedback_id,
                related_type="feedback"
            )
    
    conn.commit()
    
    return {"success": True, "message": "反馈已更新"}

@app.get("/api/v1/admin/feedback/stats")
async def admin_feedback_stats(admin: dict = Depends(get_admin_user)):
    """反馈统计（管理员仪表盘）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 总数
    cursor.execute("SELECT COUNT(*) as count FROM feedback")
    total = cursor.fetchone()["count"]
    
    # 各状态数量
    cursor.execute("SELECT status, COUNT(*) as count FROM feedback GROUP BY status")
    status_counts = {}
    for row in cursor.fetchall():
        status_counts[row["status"]] = row["count"]
    
    # 各优先级数量
    cursor.execute("SELECT priority, COUNT(*) as count FROM feedback GROUP BY priority")
    priority_counts = {}
    for row in cursor.fetchall():
        priority_counts[row["priority"]] = row["count"]
    
    # 今日新增
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE DATE(created_at) = DATE('now')")
    today_new = cursor.fetchone()["count"]
    
    pending = status_counts.get("pending", 0)
    in_progress = status_counts.get("in_progress", 0)
    resolved = status_counts.get("resolved", 0)
    
    return {
        "success": True,
        "data": {
            "total": total,
            "pending": pending,
            "in_progress": in_progress,
            "resolved": resolved,
            "today_new": today_new,
            "priority_counts": priority_counts,
            "urgent_count": priority_counts.get("urgent", 0)
        }
    }

# ==================== 通知系统 ====================

@app.get("/api/v1/notifications")
async def get_notifications(user: dict = Depends(get_current_user)):
    """获取用户通知列表"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT * FROM notifications 
        WHERE user_id = ? 
        ORDER BY created_at DESC 
        LIMIT 50
    """, (user["id"],))
    
    notifications = [dict(row) for row in cursor.fetchall()]
    
    # 未读数量
    cursor.execute("SELECT COUNT(*) as count FROM notifications WHERE user_id = ? AND is_read = 0", (user["id"],))
    unread = cursor.fetchone()["count"]
    
    conn.close()
    
    return {"success": True, "data": notifications, "unread": unread}

@app.post("/api/v1/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: int, user: dict = Depends(get_current_user)):
    """标记通知为已读"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("UPDATE notifications SET is_read = 1 WHERE id = ? AND user_id = ?", (notif_id, user["id"]))
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "已标记为已读"}

@app.post("/api/v1/notifications/read-all")
async def mark_all_notifications_read(user: dict = Depends(get_current_user)):
    """标记所有通知为已读"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("UPDATE notifications SET is_read = 1 WHERE user_id = ?", (user["id"],))
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "全部已标记为已读"}

def create_notification(user_id: int, title: str, message: str, notif_type: str = "info", related_id: int = None, related_type: str = None):
    """创建通知的辅助函数"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO notifications (user_id, title, message, type, related_id, related_type, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (user_id, title, message, notif_type, related_id, related_type, datetime.now().isoformat()))
    
    conn.commit()
    conn.close()

# ==================== 满意度评分 ====================

@app.post("/api/v1/feedback/{feedback_id}/rate")
async def rate_feedback(feedback_id: int, request: Request, user: dict = Depends(get_current_user)):
    """用户对已解决的反馈进行评分"""
    data = await request.json()
    rating = data.get("rating")
    comment = data.get("comment", "")
    
    if not rating or rating < 1 or rating > 5:
        return {"success": False, "message": "评分必须在1-5之间"}
    
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 验证反馈归属
    cursor.execute("SELECT * FROM feedback WHERE id = ? AND user_id = ?", (feedback_id, user["id"]))
    feedback = cursor.fetchone()
    
    if not feedback:
        conn.close()
        return {"success": False, "message": "反馈不存在或无权评分"}
    
    if feedback["status"] != "resolved":
        conn.close()
        return {"success": False, "message": "只能对已解决的反馈评分"}
    
    # 更新评分
    cursor.execute("""
        UPDATE feedback SET rating = ?, rating_comment = ? WHERE id = ?
    """, (rating, comment, feedback_id))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "感谢您的评分！"}

@app.get("/api/v1/admin/feedback/stats")
async def admin_feedback_stats(admin: dict = Depends(get_admin_user)):
    """反馈统计（增强版 - 管理后台仪表盘）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 总数
    cursor.execute("SELECT COUNT(*) as count FROM feedback")
    total = cursor.fetchone()["count"]
    
    # 各状态数量
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE status = 'pending'")
    pending = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE status = 'in_progress'")
    in_progress = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE status = 'resolved'")
    resolved = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE status = 'closed'")
    closed = cursor.fetchone()["count"]
    
    # 紧急数量
    cursor.execute("SELECT COUNT(*) as count FROM feedback WHERE priority = 'urgent' AND status != 'resolved' AND status != 'closed'")
    urgent_count = cursor.fetchone()["count"]
    
    # 平均满意度
    cursor.execute("SELECT AVG(rating) as avg FROM feedback WHERE rating IS NOT NULL")
    avg_rating = cursor.fetchone()["avg"] or 0
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "total": total,
            "pending": pending,
            "in_progress": in_progress,
            "resolved": resolved,
            "closed": closed,
            "urgent_count": urgent_count,
            "avg_rating": round(avg_rating, 1) if avg_rating else 0
        }
    }

@app.get("/api/v1/admin/trends")
async def admin_trends(admin: dict = Depends(get_admin_user)):
    """获取用户增长和活跃度趋势（7天数据）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    trends = {
        "user_growth": [],
        "daily_usage": [],
        "active_users": []
    }
    
    # 7天用户增长趋势
    for i in range(6, -1, -1):
        date = (datetime.now() - timedelta(days=i)).strftime("%Y-%m-%d")
        cursor.execute("""
            SELECT COUNT(*) as count FROM users 
            WHERE DATE(created_at) <= ?
        """, (date,))
        row = cursor.fetchone()
        count = row["count"] if row else 0
        trends["user_growth"].append({"date": date, "count": count})
    
    # 7天每日使用量
    for i in range(6, -1, -1):
        date = (datetime.now() - timedelta(days=i)).strftime("%Y-%m-%d")
        cursor.execute("""
            SELECT COUNT(*) as count FROM usage 
            WHERE DATE(created_at) = ?
        """, (date,))
        row = cursor.fetchone()
        count = row["count"] if row else 0
        trends["daily_usage"].append({"date": date, "count": count})
    
    # 7天活跃用户数
    for i in range(6, -1, -1):
        date = (datetime.now() - timedelta(days=i)).strftime("%Y-%m-%d")
        cursor.execute("""
            SELECT COUNT(DISTINCT user_id) as count FROM usage 
            WHERE DATE(created_at) = ?
        """, (date,))
        row = cursor.fetchone()
        count = row["count"] if row else 0
        trends["active_users"].append({"date": date, "count": count})
    
    conn.close()
    
    return {"success": True, "data": trends}

@app.get("/api/v1/admin/analytics")
async def get_user_analytics(admin: dict = Depends(get_admin_user)):
    """用户行为分析"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 活跃用户（最近7天有使用）
    cursor.execute("""
        SELECT COUNT(DISTINCT user_id) as count FROM usage 
        WHERE created_at >= DATE('now', '-7 days')
    """)
    row = cursor.fetchone()
    active_users_7d = row["count"] if row else 0
    
    # 活跃用户（最近30天有使用）
    cursor.execute("""
        SELECT COUNT(DISTINCT user_id) as count FROM usage 
        WHERE created_at >= DATE('now', '-30 days')
    """)
    row = cursor.fetchone()
    active_users_30d = row["count"] if row else 0
    
    # 用户留存率（注册后7天内仍有使用）
    cursor.execute("""
        SELECT COUNT(*) as count FROM users 
        WHERE DATE(created_at) >= DATE('now', '-7 days')
        AND id IN (SELECT DISTINCT user_id FROM usage WHERE created_at >= DATE('now', '-7 days'))
    """)
    row = cursor.fetchone()
    retained_7d = row["count"] if row else 0
    
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE DATE(created_at) >= DATE('now', '-7 days')")
    row = cursor.fetchone()
    new_7d = row["count"] if row else 0
    retention_7d = round(retained_7d / new_7d * 100, 1) if new_7d > 0 else 0
    
    # 行业使用分布
    cursor.execute("SELECT industry, COUNT(*) as count FROM usage WHERE industry IS NOT NULL GROUP BY industry ORDER BY count DESC LIMIT 10")
    industry_distribution = []
    for row in cursor.fetchall():
        industry_distribution.append({"industry": row["industry"], "count": row["count"]})
    
    # 功能使用频率
    cursor.execute("SELECT action, COUNT(*) as count FROM usage GROUP BY action ORDER BY count DESC")
    action_distribution = []
    for row in cursor.fetchall():
        action_distribution.append({"action": row["action"], "count": row["count"]})
    
    # 每日使用趋势（最近7天）
    cursor.execute("""
        SELECT DATE(created_at) as date, COUNT(*) as count 
        FROM usage WHERE created_at >= DATE('now', '-7 days')
        GROUP BY DATE(created_at) ORDER BY date
    """)
    daily_trend = []
    for row in cursor.fetchall():
        daily_trend.append({"date": row["date"], "count": row["count"]})
    
    # 用户流失率（注册后无任何使用）
    cursor.execute("""
        SELECT COUNT(*) as count FROM users 
        WHERE id NOT IN (SELECT DISTINCT user_id FROM usage)
    """)
    row = cursor.fetchone()
    inactive_users = row["count"] if row else 0
    
    # 获取总用户数
    cursor.execute("SELECT COUNT(*) as count FROM users")
    row = cursor.fetchone()
    total_users = row["count"] if row else 0
    churn_rate = round(inactive_users / total_users * 100, 1) if total_users > 0 else 0
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "active_users_7d": active_users_7d,
            "active_users_30d": active_users_30d,
            "retention_7d": retention_7d,
            "churn_rate": churn_rate,
            "inactive_users": inactive_users,
            "industry_distribution": industry_distribution,
            "action_distribution": action_distribution,
            "daily_trend": daily_trend
        }
    }

def log_error(error_type: str, error_message: str, error_stack: str = None, 
             user_id: int = None, request_path: str = None, request_method: str = None,
             severity: str = "error"):
    """记录错误日志"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO error_logs (error_type, error_message, error_stack, user_id, 
                                request_path, request_method, severity, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (error_type, error_message, error_stack, user_id, request_path, 
          request_method, severity, datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    
    logger.error(f"[{severity}] {error_type}: {error_message}")

@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    """全局异常处理"""
    import traceback
    
    log_error(
        error_type=type(exc).__name__,
        error_message=str(exc),
        error_stack=traceback.format_exc(),
        request_path=str(request.url.path),
        request_method=request.method,
        severity="critical"
    )
    
    return JSONResponse(
        status_code=500,
        content={"success": False, "detail": "服务器内部错误，请稍后重试"}
    )

@app.get("/api/v1/admin/error-logs")
async def get_error_logs(
    limit: int = 50,
    severity: Optional[str] = None,
    resolved: Optional[int] = None,
    admin: dict = Depends(get_admin_user)
):
    """获取错误日志"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    query = "SELECT * FROM error_logs WHERE 1=1"
    params = []
    
    if severity:
        query += " AND severity = ?"
        params.append(severity)
    
    if resolved is not None:
        query += " AND resolved = ?"
        params.append(resolved)
    
    query += " ORDER BY created_at DESC LIMIT ?"
    params.append(limit)
    
    cursor.execute(query, params)
    
    logs = []
    for row in cursor.fetchall():
        logs.append({
            "id": row["id"],
            "error_type": row["error_type"],
            "error_message": row["error_message"],
            "error_stack": row["error_stack"],
            "user_id": row["user_id"],
            "request_path": row["request_path"],
            "request_method": row["request_method"],
            "severity": row["severity"],
            "resolved": row["resolved"],
            "created_at": row["created_at"]
        })
    
    conn.close()
    
    return {"success": True, "data": logs, "count": len(logs)}

@app.post("/api/v1/admin/error-logs/{log_id}/resolve")
async def resolve_error_log(log_id: int, admin: dict = Depends(get_admin_user)):
    """标记错误已解决"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("UPDATE error_logs SET resolved = 1 WHERE id = ?", (log_id,))
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "错误已标记为已解决"}

@app.get("/api/v1/admin/error-stats")
async def get_error_stats(admin: dict = Depends(get_admin_user)):
    """错误统计"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 今日错误数
    cursor.execute("""
        SELECT COUNT(*) as count FROM error_logs 
        WHERE DATE(created_at) = DATE('now')
    """)
    today_errors = cursor.fetchone()["count"]
    
    # 未解决错误数
    cursor.execute("SELECT COUNT(*) as count FROM error_logs WHERE resolved = 0")
    unresolved = cursor.fetchone()["count"]
    
    # 错误类型分布
    cursor.execute("SELECT error_type, COUNT(*) as count FROM error_logs GROUP BY error_type ORDER BY count DESC LIMIT 10")
    type_distribution = []
    for row in cursor.fetchall():
        type_distribution.append({"type": row["error_type"], "count": row["count"]})
    
    # 严重程度分布
    cursor.execute("SELECT severity, COUNT(*) as count FROM error_logs GROUP BY severity")
    severity_distribution = {}
    for row in cursor.fetchall():
        severity_distribution[row["severity"]] = row["count"]
    
    conn.close()
    
    return {
        "success": True,
        "data": {
            "today_errors": today_errors,
            "unresolved": unresolved,
            "type_distribution": type_distribution,
            "severity_distribution": severity_distribution
        }
    }

@app.get("/api/v1/admin/users")
async def admin_list_users(admin: dict = Depends(get_admin_user)):
    """获取用户列表"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            u.id, u.email, u.name, u.user_level, u.level_expires_at, u.is_paid, u.is_admin, u.created_at, u.last_login,
            (SELECT COUNT(*) FROM resumes WHERE user_id = u.id) as resume_count,
            (SELECT COUNT(*) FROM usage WHERE user_id = u.id) as usage_count
        FROM users u
        ORDER BY u.created_at DESC
    """)
    
    users = []
    for row in cursor.fetchall():
        users.append({
            "id": row["id"],
            "email": row["email"],
            "name": row["name"],
            "user_level": row["user_level"] or "free",
            "level_expires_at": row["level_expires_at"],
            "level_name": USER_LEVELS.get(row["user_level"] or "free", {}).get("name", "免费用户"),
            "is_paid": bool(row["is_paid"]),
            "is_admin": bool(row["is_admin"]),
            "created_at": row["created_at"],
            "last_login": row["last_login"],
            "resume_count": row["resume_count"],
            "usage_count": row["usage_count"]
        })
    
    conn.close()
    return {"success": True, "data": users}

class UpdateUserRequest(BaseModel):
    is_paid: Optional[bool] = None
    is_admin: Optional[bool] = None
    user_level: Optional[str] = None
    level_expires_at: Optional[str] = None

@app.put("/api/v1/admin/users/{user_id}")
async def admin_update_user(user_id: int, request: UpdateUserRequest, admin: dict = Depends(get_admin_user)):
    """更新用户状态和等级"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    updates = []
    params = []
    
    if request.is_paid is not None:
        updates.append("is_paid = ?")
        params.append(1 if request.is_paid else 0)
    
    if request.is_admin is not None:
        updates.append("is_admin = ?")
        params.append(1 if request.is_admin else 0)
    
    if request.user_level is not None:
        if request.user_level not in USER_LEVELS:
            conn.close()
            raise HTTPException(status_code=400, detail="无效的用户等级")
        updates.append("user_level = ?")
        params.append(request.user_level)
    
    if request.level_expires_at is not None:
        updates.append("level_expires_at = ?")
        params.append(request.level_expires_at)
    
    if not updates:
        conn.close()
        raise HTTPException(status_code=400, detail="没有要更新的字段")
    
    params.append(user_id)
    cursor.execute(f"UPDATE users SET {', '.join(updates)} WHERE id = ?", params)
    conn.commit()
    conn.close()
    
    return {"success": True, "message": "用户已更新"}

@app.get("/api/v1/admin/usage")
async def admin_usage_logs(
    email: Optional[str] = None,
    action: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    admin: dict = Depends(get_admin_user)
):
    """获取使用日志（支持筛选）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 构建查询
    query = """
        SELECT 
            u.id, u.user_id, u.action, u.industry, u.created_at,
            us.email, us.name as user_name
        FROM usage u
        LEFT JOIN users us ON u.user_id = us.id
        WHERE 1=1
    """
    params = []
    
    if email:
        query += " AND us.email LIKE ?"
        params.append(f"%{email}%")
    
    if action:
        query += " AND u.action = ?"
        params.append(action)
    
    if date_from:
        query += " AND u.created_at >= ?"
        params.append(f"{date_from}T00:00:00")
    
    if date_to:
        query += " AND u.created_at <= ?"
        params.append(f"{date_to}T23:59:59")
    
    query += " ORDER BY u.created_at DESC LIMIT 200"
    
    cursor.execute(query, params)
    
    logs = []
    for row in cursor.fetchall():
        logs.append({
            "id": row["id"],
            "user_id": row["user_id"],
            "email": row["email"] or "-",
            "user_name": row["user_name"] or "-",
            "action": row["action"],
            "industry": row["industry"],
            "created_at": row["created_at"]
        })
    
    conn.close()
    return {"success": True, "data": logs, "count": len(logs)}

class TestAIRequest(BaseModel):
    provider: str = "dashscope"
    model: str = "qwen-turbo"
    api_key: Optional[str] = None
    base_url: Optional[str] = None

class MarketingEmailRequest(BaseModel):
    email: str
    template_type: str  # 'welcome', 'usage_reminder', 'upgrade_reminder', 'promo'
    data: Optional[dict] = {}

class BatchMarketingEmailRequest(BaseModel):
    target: str  # 'inactive', 'all'
    template_type: str

@app.post("/api/v1/admin/test-ai")
async def admin_test_ai(request: TestAIRequest, admin: dict = Depends(get_admin_user)):
    """测试AI连接"""
    import httpx
    
    # 使用请求中的配置或默认配置
    api_key = request.api_key or DASHSCOPE_API_KEY
    base_url = request.base_url or DASHSCOPE_BASE_URL
    model = request.model or DASHSCOPE_MODEL
    
    if not api_key:
        return {"success": False, "error": "未配置API Key"}
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "User-Agent": "Mozilla/5.0 ResumeAI/1.0"  # Coding API需要User-Agent
                },
                json={
                    "model": model,
                    "messages": [
                        {"role": "user", "content": "请回复：AI连接测试成功"}
                    ],
                    "max_tokens": 50
                }
            )
            
            if response.status_code == 200:
                data = response.json()
                reply = data.get("choices", [{}])[0].get("message", {}).get("content", "")
                return {"success": True, "response": reply, "model": model}
            else:
                error_msg = response.text[:200]
                return {"success": False, "error": f"API错误: {response.status_code} - {error_msg}"}
    except Exception as e:
        return {"success": False, "error": f"连接失败: {str(e)}"}

@app.post("/api/v1/admin/send-marketing-email")
async def admin_send_marketing_email(
    request: MarketingEmailRequest,
    admin: dict = Depends(get_admin_user)
):
    """发送营销邮件"""
    result = send_marketing_email(request.email, request.template_type, request.data)
    return {"success": result["success"], "message": result["message"]}

@app.post("/api/v1/admin/batch-marketing-email")
async def admin_batch_marketing_email(
    request: BatchMarketingEmailRequest,
    admin: dict = Depends(get_admin_user)
):
    """批量发送营销邮件"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取目标用户
    if request.target == "inactive":
        # 获取无使用记录的用户
        cursor.execute("""
            SELECT email FROM users 
            WHERE id NOT IN (SELECT DISTINCT user_id FROM usage)
            AND email != 'zhwffy@hotmail.com'
            LIMIT 50
        """)
    elif request.target == "all":
        cursor.execute("SELECT email FROM users WHERE email != 'zhwffy@hotmail.com' LIMIT 50")
    else:
        return {"success": False, "message": "未知目标群体"}
    
    users = cursor.fetchall()
    conn.close()
    
    sent_count = 0
    failed_count = 0
    
    for user in users:
        result = send_marketing_email(user["email"], request.template_type)
        if result["success"]:
            sent_count += 1
        else:
            failed_count += 1
    
    return {
        "success": True,
        "sent_count": sent_count,
        "failed_count": failed_count,
        "total": len(users)
    }

@app.post("/api/v1/admin/backup")
async def create_backup(admin: dict = Depends(get_admin_user)):
    """创建数据库备份"""
    import shutil
    import os
    
    backup_dir = "backups"
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_filename = f"{backup_dir}/resume_ai_backup_{timestamp}.db"
    
    # 备份数据库文件
    source_db = "resume_users.db"
    if os.path.exists(source_db):
        shutil.copy2(source_db, backup_filename)
        
        return {
            "success": True,
            "message": "备份成功",
            "backup_file": backup_filename,
            "backup_time": timestamp
        }
    else:
        return {"success": False, "message": "数据库文件不存在"}

@app.get("/api/v1/admin/backups")
async def list_backups(admin: dict = Depends(get_admin_user)):
    """获取备份列表"""
    import os
    
    backup_dir = "backups"
    if not os.path.exists(backup_dir):
        return {"success": True, "data": [], "message": "备份目录不存在"}
    
    backups = []
    for filename in os.listdir(backup_dir):
        if filename.endswith(".db"):
            filepath = os.path.join(backup_dir, filename)
            file_size = os.path.getsize(filepath)
            file_time = datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
            
            backups.append({
                "filename": filename,
                "size": file_size,
                "created_at": file_time,
                "path": filepath
            })
    
    # 按时间倒序
    backups.sort(key=lambda x: x["created_at"], reverse=True)
    
    return {"success": True, "data": backups, "count": len(backups)}

@app.post("/api/v1/admin/restore/{backup_filename}")
async def restore_backup(backup_filename: str, admin: dict = Depends(get_admin_user)):
    """从备份恢复数据库"""
    import shutil
    import os
    
    backup_dir = "backups"
    backup_path = os.path.join(backup_dir, backup_filename)
    
    if not os.path.exists(backup_path):
        raise HTTPException(status_code=404, detail="备份文件不存在")
    
    # 恢复数据库
    source_db = "resume_users.db"
    
    # 先备份当前数据库
    if os.path.exists(source_db):
        current_backup = f"{backup_dir}/pre_restore_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
        shutil.copy2(source_db, current_backup)
    
    # 恢复备份
    shutil.copy2(backup_path, source_db)
    
    return {
        "success": True,
        "message": "恢复成功",
        "restored_from": backup_filename,
        "pre_restore_backup": current_backup if os.path.exists(source_db) else None
    }

@app.get("/api/v1/admin/config")
async def admin_get_config(admin: dict = Depends(get_admin_user)):
    """获取系统配置"""
    config = {
        "email": {
            "smtp_server": EMAIL_CONFIG.get("smtp_server", ""),
            "smtp_port": EMAIL_CONFIG.get("smtp_port", 465),
            "smtp_user": EMAIL_CONFIG.get("smtp_user", ""),
            "configured": EMAIL_CONFIG.get("smtp_configured", False),
            "test_mode": EMAIL_CONFIG.get("test_mode", True)
        },
        "ai": {
            "provider": "coding",
            "model": DASHSCOPE_MODEL,
            "base_url": DASHSCOPE_BASE_URL,
            "configured": bool(DASHSCOPE_API_KEY)
        },
        "levels": USER_LEVELS
    }
    return {"success": True, "data": config}

if __name__ == "__main__":
    import uvicorn
    print(f"🚀 ResumeAI Backend Starting...")
    print(f"Keywords loaded: {len(KEYWORDS_DB)} industries")
    print(f"Auth enabled: True")
    print(f"Free limit: {FREE_LIMIT} times/day")
    uvicorn.run(app, host="0.0.0.0", port=8001)

# Vercel Serverless Handler
try:
    from mangum import Mangum
    handler = Mangum(app, lifespan="off")
except ImportError:
    # 如果mangum未安装，提供一个基本的handler
    def handler(event, context):
        return {
            "statusCode": 500,
            "body": "Mangum not installed",
            "headers": {"Content-Type": "text/plain"}
        }