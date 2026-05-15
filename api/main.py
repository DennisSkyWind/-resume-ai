"""
ResumeAI Backend - FastAPIwith User Authentication
简历优化系统后端服务（带用户认证）
"""

from fastapi import FastAPI, HTTPException, Depends, Header, File, UploadFile, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List
import json
import os
import re
import sqlite3
import hashlib
import secrets
from datetime import datetime, timedelta
import jwt
import random
import tempfile
import logging

# 配置日志 - Vercel兼容（无文件日志）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 导入邮箱发送模块
from email_sender import send_verification_code, generate_code, get_email_config, EMAIL_CONFIG

# 文件解析模块
import pdfplumber
from docx import Document

# PDF生成模块
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO

# 初始化FastAPI
app = FastAPI(
    title="ResumeAI API",
    description="AI简历优化系统 - 为非IT行业求职者提供简历优化服务",
    version="2.0.0"
)

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 数据目录 - 多环境兼容
# Render: 使用backend目录下的data子目录
# Vercel: 使用/tmp目录（Serverless）
# 本地: 使用项目根目录的data目录
def get_data_dir():
    # 检查环境变量
    if os.environ.get("VERCEL") == "1":
        return "/tmp"
    
    # Render或本地环境
    # 尝试多个可能的路径
    possible_paths = [
        os.path.join(os.path.dirname(__file__), "data"),  # backend/data
        os.path.join(os.path.dirname(__file__), "..", "data"),  # project_root/data
        "/tmp"  # fallback
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
        # 尝试创建目录
        try:
            os.makedirs(path, exist_ok=True)
            return path
        except:
            continue
    
    return "/tmp"

DATA_DIR = get_data_dir()
USER_DB_PATH = os.path.join(DATA_DIR, "users.db")

# ========== 性能优化：内存缓存 ==========
from functools import lru_cache
import time

# 缓存配置
CACHE_TTL_SECONDS = 300  # 缓存有效期5分钟

# 简单缓存装饰器（带TTL）
def cached_with_ttl(ttl_seconds=CACHE_TTL_SECONDS):
    """带过期时间的缓存装饰器"""
    cache = {}
    
    def decorator(func):
        def wrapper(*args, **kwargs):
            key = str(args) + str(kwargs)
            now = time.time()
            
            if key in cache:
                value, timestamp = cache[key]
                if now - timestamp < ttl_seconds:
                    return value
            
            result = func(*args, **kwargs)
            cache[key] = (result, now)
            return result
        
        wrapper.cache_clear = lambda: cache.clear()
        return wrapper
    
    return decorator

# 统计数据缓存（5分钟刷新）
_stats_cache = {"data": None, "timestamp": 0}

def get_cached_stats():
    """获取缓存的统计数据"""
    global _stats_cache
    now = time.time()
    
    if _stats_cache["data"] and (now - _stats_cache["timestamp"] < CACHE_TTL_SECONDS):
        return _stats_cache["data"]
    
    return None

def set_cached_stats(data):
    """更新统计数据缓存"""
    global _stats_cache
    _stats_cache = {"data": data, "timestamp": time.time()}

# 确保数据库文件存在
if not os.path.exists(USER_DB_PATH):
    try:
        conn = sqlite3.connect(USER_DB_PATH)
        # 用户表
        conn.execute('''CREATE TABLE IF NOT EXISTS users
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE,
            password_hash TEXT,
            name TEXT,
            verified INTEGER DEFAULT 0,
            daily_limit INTEGER DEFAULT 5,
            last_reset TEXT,
            last_login TEXT,
            created_at TEXT)''')
        # 验证码表
        conn.execute('''CREATE TABLE IF NOT EXISTS verification_codes
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT,
            code TEXT,
            expires_at TEXT,
            used INTEGER DEFAULT 0,
            created_at TEXT)''')
        conn.commit()
        
        # conn.close() # 优化：使用共享连接
        logger.info(f"Created database at {USER_DB_PATH}")
    except Exception as e:
        logger.error(f"Failed to create database: {e}")

# 性能优化：应用启动时创建索引
@app.on_event("startup")
async def create_indexes():
    """启动时创建数据库索引以提升查询性能"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 检查并创建索引
    indexes = [
        ("idx_users_email", "users(email)"),
        ("idx_users_created_at", "users(created_at)"),
        ("idx_usage_user_id", "usage(user_id)"),
        ("idx_usage_created_at", "usage(created_at)"),
        ("idx_resumes_user_id", "resumes(user_id)"),
        ("idx_error_logs_created_at", "error_logs(created_at)"),
    ]
    
    for idx_name, idx_target in indexes:
        try:
            cursor.execute(f"CREATE INDEX IF NOT EXISTS {idx_name} ON {idx_target}")
        except sqlite3.OperationalError:
            pass  # 表不存在，跳过
    
    conn.commit()
    logger.info("Database indexes created/verified")

# JWT配置（使用固定密钥）
JWT_SECRET = "resumeai_jwt_secret_key_2026"  # 固定密钥，生产环境应使用环境变量
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 168  # 7天有效期

# 免费用户限制
FREE_LIMIT = 5  # 每天免费使用次数

# IP速率限制配置
RATE_LIMIT_WINDOW = 60  # 60秒窗口
RATE_LIMIT_MAX_REQUESTS = 20  # 每分钟最大请求次数
ip_request_counts = {}  # IP请求计数（内存存储）

def check_rate_limit(ip: str) -> bool:
    """检查IP速率限制"""
    now = datetime.now()
    window_start = now - timedelta(seconds=RATE_LIMIT_WINDOW)
    
    # 清理过期记录
    if ip in ip_request_counts:
        ip_request_counts[ip] = [t for t in ip_request_counts[ip] if t > window_start]
    else:
        ip_request_counts[ip] = []
    
    # 检查请求次数
    if len(ip_request_counts[ip]) >= RATE_LIMIT_MAX_REQUESTS:
        return False
    
    # 记录本次请求
    ip_request_counts[ip].append(now)
    return True

# PDF中文字体配置
FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
FONT_BOLD_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
try:
    pdfmetrics.registerFont(TTFont('NotoSansCJK', FONT_PATH, subfontIndex=0))
    pdfmetrics.registerFont(TTFont('NotoSansCJK-Bold', FONT_BOLD_PATH, subfontIndex=0))
    PDF_FONT_NAME = 'NotoSansCJK'
    PDF_FONT_BOLD = 'NotoSansCJK-Bold'
except Exception as e:
    print(f"字体注册失败: {e}")
    PDF_FONT_NAME = 'Helvetica'
    PDF_FONT_BOLD = 'Helvetica-Bold'

# 阿里云Coding API配置
DASHSCOPE_API_KEY = "sk-sp-e8d1076e8dd4461d8d1edf2542f8de68"
DASHSCOPE_BASE_URL = "https://coding.dashscope.aliyuncs.com/v1"
DASHSCOPE_MODEL = "qwen3.5-plus"

# 加载关键词库
def load_keywords():
    # Vercel环境：从api目录读取
    keywords_file = os.path.join(os.path.dirname(__file__), "keywords.json")
    if os.path.exists(keywords_file):
        with open(keywords_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    # 兜底：尝试data目录
    keywords_file = os.path.join(DATA_DIR, "keywords.json")
    if os.path.exists(keywords_file):
        with open(keywords_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

KEYWORDS_DB = load_keywords()

# ========== 用户认证函数 ==========

# 性能优化：数据库连接缓存和SQLite调优
_db_connection = None

def get_user_db():
    """获取数据库连接（性能优化版）"""
    global _db_connection
    if _db_connection is None:
        conn = sqlite3.connect(USER_DB_PATH, check_same_thread=False)
        conn.row_factory = sqlite3.Row
        # SQLite性能优化PRAGMA
        conn.execute("PRAGMA journal_mode=WAL")  # 写前日志，提升并发性能
        conn.execute("PRAGMA synchronous=NORMAL")  # 降低同步频率
        conn.execute("PRAGMA cache_size=-64000")  # 64MB缓存
        conn.execute("PRAGMA temp_store=MEMORY")  # 临时表存内存
        _db_connection = conn
    return _db_connection

def close_db_connection():
    """关闭数据库连接（服务关闭时调用）"""
    global _db_connection
    if _db_connection:
        _db_connection.close()
        _db_connection = None

@app.on_event("shutdown")
async def shutdown_event():
    close_db_connection()

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def check_password_strength(password: str) -> dict:
    """检查密码强度"""
    result = {
        "valid": True,
        "score": 0,
        "message": ""
    }
    
    # 长度检查
    if len(password) < 6:
        result["valid"] = False
        result["message"] = "密码长度至少6位"
        return result
    elif len(password) >= 8:
        result["score"] += 1
    
    # 复杂度检查
    has_lower = any(c.islower() for c in password)
    has_upper = any(c.isupper() for c in password)
    has_digit = any(c.isdigit() for c in password)
    has_special = any(c in "!@#$%^&*()_+-=[]{}|;:,.<>?/" for c in password)
    
    complexity = sum([has_lower, has_upper, has_digit, has_special])
    result["score"] += complexity
    
    # 评分
    if result["score"] >= 4:
        result["message"] = "密码强度：强"
    elif result["score"] >= 2:
        result["message"] = "密码强度：中等"
    else:
        result["message"] = "密码强度：弱，建议包含大小写字母、数字和特殊字符"
    
    return result

def create_token(user_id: int, email: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def verify_token(token: str) -> Optional[dict]:
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def detect_language(text: str) -> str:
    """检测文本语言（中文/英文）"""
    # 统计中文字符比例
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(text.strip())
    
    if total_chars == 0:
        return "zh"  # 默认中文
    
    chinese_ratio = chinese_chars / total_chars
    
    # 中文字符占比超过30%视为中文
    if chinese_ratio > 0.3:
        return "zh"
    else:
        return "en"

async def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="未提供认证token")
    
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    payload = verify_token(token)
    
    if not payload:
        raise HTTPException(status_code=401, detail="token无效或已过期")
    
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (payload["user_id"],))
    user = cursor.fetchone()
    # conn.close() # 优化：使用共享连接
    
    if not user:
        raise HTTPException(status_code=401, detail="用户不存在")
    
    return dict(user)

def check_usage_limit(user_id: int) -> bool:
    """检查用户今日使用次数"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 检查是否付费用户
    cursor.execute("SELECT is_paid FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()
    
    if user and user["is_paid"]:
        # conn.close() # 优化：使用共享连接
        return True  # 付费用户无限制
    
    # 检查今日使用次数
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user_id, today))
    
    result = cursor.fetchone()
    # conn.close() # 优化：使用共享连接
    
    return result["count"] < FREE_LIMIT

def record_usage(user_id: int, action: str, industry: str = None):
    """记录用户使用"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO usage (user_id, action, industry)
        VALUES (?, ?, ?)
    """, (user_id, action, industry))
    conn.commit()
    # conn.close() # 优化：使用共享连接

# ========== 请求模型 ==========

class SendCodeRequest(BaseModel):
    email: str

class RegisterRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = None
    code: str  # 验证码

class LoginRequest(BaseModel):
    email: str
    password: str

class AnalyzeRequest(BaseModel):
    resume_content: str
    industry: str
    language: str = "zh"
    jd_content: Optional[str] = None

class OptimizeRequest(BaseModel):
    resume_content: str
    industry: str
    language: str = "zh"
    jd_content: Optional[str] = None

# ========== 用户认证API ==========

@app.post("/api/v1/send-code")
async def send_code(request: SendCodeRequest):
    """发送验证码"""
    try:
        import sqlite3
        
        conn = get_user_db()
        cursor = conn.cursor()
        
        # 检查邮箱是否已注册
        cursor.execute("SELECT id FROM users WHERE email = ?", (request.email,))
        if cursor.fetchone():
            # conn.close() # 优化：使用共享连接
            raise HTTPException(status_code=400, detail="邮箱已被注册")
        
        # 生成验证码
        code = generate_code(EMAIL_CONFIG["code_length"])
        expires_at = datetime.now() + timedelta(minutes=EMAIL_CONFIG["code_expire_minutes"])
        
        # 保存验证码到数据库
        cursor.execute("""
            INSERT INTO verification_codes (email, code, expires_at)
            VALUES (?, ?, ?)
        """, (request.email, code, expires_at.isoformat()))
        conn.commit()
        # conn.close() # 优化：使用共享连接
        
        # 发送验证码
        result = send_verification_code(request.email, code)
        
        return {
            "success": result["success"],
            "message": result["message"],
            "test_code": result.get("test_code")  # 测试模式返回验证码
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"发送验证码失败: {e}")
        raise HTTPException(status_code=500, detail=f"发送失败: {str(e)}")

@app.post("/api/v1/register")
async def register(request: RegisterRequest):
    """用户注册（需要验证码）"""
    logger.info(f"注册请求: email={request.email}")
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 验证邮箱格式
    if not re.match(r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$", request.email):
        logger.warning(f"注册失败: {request.email} - 邮箱格式不正确")
        raise HTTPException(status_code=400, detail="邮箱格式不正确")
    
    # 检查密码强度
    strength = check_password_strength(request.password)
    if not strength["valid"]:
        logger.warning(f"注册失败: {request.email} - 密码强度不足")
        raise HTTPException(status_code=400, detail=strength["message"])
    
    # 检查邮箱是否已存在
    cursor.execute("SELECT id FROM users WHERE email = ?", (request.email,))
    if cursor.fetchone():
        # conn.close() # 优化：使用共享连接
        logger.warning(f"注册失败: {request.email} - 邮箱已被注册")
        raise HTTPException(status_code=400, detail="邮箱已被注册")
    
    # 验证验证码
    cursor.execute("""
        SELECT * FROM verification_codes 
        WHERE email = ? AND code = ? AND used = FALSE AND expires_at > ?
        ORDER BY created_at DESC LIMIT 1
    """, (request.email, request.code, datetime.now().isoformat()))
    
    code_record = cursor.fetchone()
    
    if not code_record:
        # conn.close() # 优化：使用共享连接
        raise HTTPException(status_code=400, detail="验证码无效或已过期")
    
    # 标记验证码已使用
    cursor.execute("UPDATE verification_codes SET used = TRUE WHERE id = ?", (code_record["id"],))
    
    # 创建用户
    password_hash = hash_password(request.password)
    cursor.execute("""
        INSERT INTO users (email, password_hash, name)
        VALUES (?, ?, ?)
    """, (request.email, password_hash, request.name))
    
    user_id = cursor.lastrowid
    conn.commit()
    # conn.close() # 优化：使用共享连接
    
    # 生成token
    token = create_token(user_id, request.email)
    
    return {
        "success": True,
        "data": {
            "user_id": user_id,
            "email": request.email,
            "token": token,
            "token_type": "Bearer"
        }
    }

@app.post("/api/v1/login")
async def login(request: LoginRequest):
    """用户登录"""
    logger.info(f"登录请求: {request.email}")
    conn = get_user_db()
    cursor = conn.cursor()
    
    password_hash = hash_password(request.password)
    cursor.execute("""
            SELECT id, email, name, daily_limit FROM users WHERE email = ? AND password_hash = ?
        """, (request.email, password_hash))
    
    user = cursor.fetchone()
    
    if not user:
        logger.warning(f"登录失败: {request.email} - 邮箱或密码错误")
        # conn.close() # 优化：使用共享连接
        raise HTTPException(status_code=401, detail="邮箱或密码错误")
    
    # 更新登录时间
    cursor.execute("""
        UPDATE users SET last_login = ? WHERE id = ?
    """, (datetime.now().isoformat(), user["id"]))
    conn.commit()
    # conn.close() # 优化：使用共享连接
    
    # 生成token
    token = create_token(user["id"], user["email"])
    
    return {
        "success": True,
        "data": {
            "user_id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "daily_limit": user["daily_limit"],
            "token": token,
            "token_type": "Bearer"
        }
    }

@app.get("/api/v1/me")
async def get_me(user: dict = Depends(get_current_user)):
    """获取当前用户信息"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取今日使用次数
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user["id"], today))
    
    usage = cursor.fetchone()
    
    # 获取简历数量
    cursor.execute("SELECT COUNT(*) as count FROM resumes WHERE user_id = ?", (user["id"],))
    resumes = cursor.fetchone()
    # conn.close() # 优化：使用共享连接
    
    return {
        "success": True,
        "data": {
            "user_id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "is_admin": bool(user.get("is_admin", 0)),
            "is_paid": user["is_paid"],
            "plan_type": user["plan_type"],
            "usage_today": usage["count"],
            "usage_limit": FREE_LIMIT if not user["is_paid"] else "unlimited",
            "resumes_count": resumes["count"]
        }
    }

@app.get("/api/v1/usage")
async def get_usage(user: dict = Depends(get_current_user)):
    """获取用户使用次数详情"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取今日使用次数
    today = datetime.now().strftime("%Y-%m-%d")
    cursor.execute("""
        SELECT COUNT(*) as count FROM usage 
        WHERE user_id = ? AND DATE(created_at) = ?
    """, (user["id"], today))
    
    usage = cursor.fetchone()
    used_today = usage["count"]
    
    # 获取是否付费
    is_paid = user["is_paid"]
    remaining = "unlimited" if is_paid else max(0, FREE_LIMIT - used_today)
    
    # conn.close() # 优化：使用共享连接
    
    return {
        "success": True,
        "data": {
            "used_today": used_today,
            "limit": FREE_LIMIT if not is_paid else "unlimited",
            "remaining": remaining,
            "is_paid": is_paid,
            "can_use": is_paid or used_today < FREE_LIMIT
        }
    }

# ========== 简历API（需认证） ==========

@app.post("/api/v1/upload")
async def upload_resume(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    """上传简历文件并提取文本（PDF/Word）"""
    
    # 检查文件类型
    filename = file.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.docx') or filename.endswith('.doc')):
        raise HTTPException(status_code=400, detail="只支持PDF和Word文件（.pdf, .docx, .doc），请将图片或截图转换为PDF后上传")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完")
    
    try:
        # 保存临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        text_content = ""
        
        # 解析PDF
        if filename.endswith('.pdf'):
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        
        # 解析Word
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text_content += " | ".join(row_text) + "\n"
        
        # 清理临时文件
        os.unlink(tmp_path)
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="无法从文件中提取文本内容")
        
        # 记录使用
        record_usage(user["id"], "upload", "file")
        
        return {
            "success": True,
            "data": {
                "filename": file.filename,
                "content": text_content.strip(),
                "length": len(text_content.strip())
            }
        }
        
    except Exception as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

@app.post("/api/v1/analyze")
async def analyze(request: AnalyzeRequest, user: dict = Depends(get_current_user), 
                  client_ip: Optional[str] = Header(None, alias="X-Forwarded-For")):
    """分析简历（需认证）"""
    # IP速率限制
    ip = client_ip or "unknown"
    if not check_rate_limit(ip):
        logger.warning(f"速率限制触发: IP={ip}")
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    
    logger.info(f"简历分析请求: user={user['id']}, industry={request.industry}, ip={ip}")
    
    if request.industry not in KEYWORDS_DB:
        raise HTTPException(status_code=400, detail=f"不支持的行业: {request.industry}")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完，请升级付费或明天再试")
    
    # 分析简历
    result = analyze_resume(request.resume_content, request.industry, request.language, request.jd_content)
    
    # 记录使用
    record_usage(user["id"], "analyze", request.industry)
    
    # 保存简历
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO resumes (user_id, content, industry, score)
        VALUES (?, ?, ?, ?)
    """, (user["id"], request.resume_content, request.industry, result["overall_score"]))
    conn.commit()
    # conn.close() # 优化：使用共享连接
    
    return {"success": True, "data": result}

@app.post("/api/v1/optimize")
async def optimize(request: OptimizeRequest, user: dict = Depends(get_current_user)):
    """优化简历（需认证）"""
    if request.industry not in KEYWORDS_DB:
        raise HTTPException(status_code=400, detail=f"不支持的行业: {request.industry}")
    
    # 检查使用限制
    if not check_usage_limit(user["id"]):
        raise HTTPException(status_code=403, detail="今日免费次数已用完")
    
    # 优化简历（简化版）
    result = optimize_resume(request.resume_content, request.industry, request.language)
    
    # 记录使用
    record_usage(user["id"], "optimize", request.industry)
    
    return {"success": True, "data": result}

@app.get("/api/v1/resumes")
async def get_resumes(user: dict = Depends(get_current_user)):
    """获取用户简历历史"""
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, title, industry, score, created_at
        FROM resumes WHERE user_id = ?
        ORDER BY created_at DESC LIMIT 20
    """, (user["id"],))
    
    resumes = [dict(row) for row in cursor.fetchall()]
    # conn.close() # 优化：使用共享连接
    
    return {"success": True, "data": resumes}

@app.get("/api/v1/resumes/history")
async def get_resumes_history(user: dict = Depends(get_current_user)):
    """获取用户简历历史列表（带分页）"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 获取总数
    cursor.execute("SELECT COUNT(*) as total FROM resumes WHERE user_id = ?", (user["id"],))
    total = cursor.fetchone()["total"]
    
    # 获取历史列表（带更多字段）
    cursor.execute("""
        SELECT 
            id,
            title,
            industry,
            score,
            created_at,
            updated_at,
            LENGTH(content) as content_length
        FROM resumes 
        WHERE user_id = ?
        ORDER BY created_at DESC
        LIMIT 50
    """, (user["id"],))
    
    history = [dict(row) for row in cursor.fetchall()]
    # conn.close() # 优化：使用共享连接
    
    return {
        "success": True,
        "data": {
            "total": total,
            "history": history
        }
    }

@app.get("/api/v1/resumes/{resume_id}")
async def get_resume_detail(resume_id: int, user: dict = Depends(get_current_user)):
    """获取简历详情"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            id,
            user_id,
            title,
            industry,
            score,
            content,
            created_at,
            updated_at
        FROM resumes 
        WHERE id = ? AND user_id = ?
    """, (resume_id, user["id"]))
    
    resume = cursor.fetchone()
    # conn.close() # 优化：使用共享连接
    
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在或无权访问")
    
    return {
        "success": True,
        "data": dict(resume)
    }

@app.post("/api/v1/export/pdf")
async def export_pdf(user: dict = Depends(get_current_user)):
    """导出简历分析报告PDF"""
    # 生成PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    
    # 样式配置
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        fontName=PDF_FONT_BOLD,
        fontSize=18,
        spaceAfter=20,
        alignment=1  # 居中
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        fontName=PDF_FONT_BOLD,
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'CustomBody',
        fontName=PDF_FONT_NAME,
        fontSize=12,
        spaceBefore=6,
        spaceAfter=6,
        leading=18
    )
    
    # 构建PDF内容
    elements = []
    
    # 标题
    elements.append(Paragraph("简历分析报告", title_style))
    elements.append(Spacer(1, 20))
    
    # 基本信息
    elements.append(Paragraph("基本信息", heading_style))
    info_data = [
        ["分析日期", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["用户邮箱", user.get("email", "未知")],
    ]
    info_table = Table(info_data, colWidths=[80*mm, 80*mm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), PDF_FONT_NAME),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 30))
    
    # 落款
    elements.append(Spacer(1, 50))
    elements.append(Paragraph(
        f"{datetime.now().strftime('%Y年%m月%d日')} 智能与数字化中心·数据智能",
        ParagraphStyle('Footer', fontName=PDF_FONT_NAME, fontSize=12, alignment=2)
    ))
    
    # 生成PDF
    doc.build(elements)
    buffer.seek(0)
    
    # 返回PDF响应
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=resume_analysis_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

# ========== 核心分析函数 ==========

def analyze_resume(resume: str, industry: str, language: str, jd: Optional[str] = None):
    """分析简历关键词匹配度"""
    
    industry_keywords = KEYWORDS_DB.get(industry, {})
    required_keywords = industry_keywords.get("keywords", {}).get("required", [])
    preferred_keywords = industry_keywords.get("keywords", {}).get("preferred", [])
    
    all_keywords = required_keywords + preferred_keywords
    
    keywords_found = []
    keywords_missing = []
    total_weight = 0
    matched_weight = 0
    
    for kw in all_keywords:
        keyword_text = kw["keyword"] if language == "zh" else kw.get("en", kw["keyword"])
        count = len(re.findall(keyword_text, resume, re.IGNORECASE))
        found = count > 0
        
        if found:
            keywords_found.append({
                "keyword": kw["keyword"],
                "weight": kw["weight"],
                "count": count,
                "found": True
            })
            matched_weight += kw["weight"]
        
        total_weight += kw["weight"]
    
    # 计算缺失关键词
    found_kw_list = [k["keyword"] for k in keywords_found]
    for kw in all_keywords:
        if kw["keyword"] not in found_kw_list:
            keywords_missing.append({
                "keyword": kw["keyword"],
                "weight": kw["weight"],
                "suggestion": f"添加{kw['keyword']}相关经验描述"
            })
    
    # 计算评分
    keyword_score = round(matched_weight / total_weight * 100, 1) if total_weight > 0 else 0
    ats_score = check_ats(resume)
    overall_score = round((keyword_score * 0.6 + ats_score * 0.4), 1)
    
    suggestions = generate_suggestions(resume, keywords_missing, industry)
    
    return {
        "overall_score": overall_score,
        "keyword_score": keyword_score,
        "ats_score": ats_score,
        "keywords_found": keywords_found,
        "keywords_missing": keywords_missing,
        "suggestions": suggestions,
        "industry_name": industry_keywords.get("name", industry)
    }

def optimize_resume(resume: str, industry: str, language: str):
    """优化简历"""
    analysis = analyze_resume(resume, industry, language)
    
    optimized = resume
    
    # 添加建议
    suggestions_text = "\n\n".join([s["suggestion"] for s in analysis["suggestions"]])
    
    return {
        "optimized_resume": optimized,
        "added_suggestions": suggestions_text,
        "score_improvement": f"建议添加{len(analysis['keywords_missing'])}个关键词",
        "analysis": analysis
    }

def check_ats(resume: str):
    """ATS兼容性检测"""
    score = 100
    
    if len(resume) < 200:
        score -= 20
    if len(resume) > 5000:
        score -= 10
    if not re.search(r'\d{11}', resume) and not re.search(r'[\w\.-]+@[\w\.-]+\.\w+', resume):
        score -= 15
    if re.search(r'[◆●■□▲►]', resume):
        score -= 5
    
    return max(0, score)

def generate_suggestions(resume: str, missing_keywords: list, industry: str):
    """生成优化建议"""
    suggestions = []
    industry_data = KEYWORDS_DB.get(industry, {})
    
    for kw in missing_keywords[:5]:
        suggestions.append({
            "type": "add_keyword",
            "keyword": kw["keyword"],
            "suggestion": f"建议添加「{kw['keyword']}」相关经验描述"
        })
    
    if not re.search(r'\d+%', resume) and not re.search(r'\d+万', resume):
        metrics = industry_data.get("metrics_examples", [])
        if metrics:
            suggestions.append({
                "type": "add_metric",
                "suggestion": f"建议添加量化数据，如：{metrics[0]}"
            })
    
    return suggestions

# ========== 其他API ==========

@app.get("/")
async def root():
    return {
        "name": "ResumeAI API",
        "version": "2.0.0",
        "industries": list(KEYWORDS_DB.keys()),
        "auth_enabled": True
    }

@app.get("/api/v1/industries")
async def get_industries():
    industries = []
    for key, data in KEYWORDS_DB.items():
        industries.append({
            "id": key,
            "name": data.get("name", key),
            "name_en": data.get("name_en", key)
        })
    return {"success": True, "data": industries}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "keywords_loaded": len(KEYWORDS_DB),
        "auth_enabled": True,
        "data_dir": DATA_DIR,
        "db_path": USER_DB_PATH
    }

@app.get("/debug/db")
async def debug_db():
    """调试数据库状态"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        # 检查表
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        # 检查users表结构
        cursor.execute("PRAGMA table_info(users)")
        users_columns = [row[1] for row in cursor.fetchall()]
        # 检查用户数量
        cursor.execute("SELECT COUNT(*) FROM users")
        user_count = cursor.fetchone()[0]
        # conn.close() # 优化：使用共享连接
        return {
            "tables": tables,
            "users_columns": users_columns,
            "user_count": user_count,
            "db_path": USER_DB_PATH,
            "data_dir": DATA_DIR
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/debug/init-tables")
async def init_tables():
    """初始化/修复数据库表结构"""
    try:
        conn = sqlite3.connect(USER_DB_PATH)
        cursor = conn.cursor()
        
        # 检查users表结构
        cursor.execute("PRAGMA table_info(users)")
        columns = [row[1] for row in cursor.fetchall()]
        
        # 添加缺失的字段
        missing_columns = []
        
        if 'is_admin' not in columns:
            cursor.execute("ALTER TABLE users ADD COLUMN is_admin INTEGER DEFAULT 0")
            missing_columns.append('is_admin')
        
        if 'is_paid' not in columns:
            cursor.execute("ALTER TABLE users ADD COLUMN is_paid INTEGER DEFAULT 0")
            missing_columns.append('is_paid')
        
        # 创建verification_codes表
        cursor.execute('''CREATE TABLE IF NOT EXISTS verification_codes
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT,
            code TEXT,
            expires_at TEXT,
            used INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP)''')
        
        conn.commit()
        
        # 检查最终表结构
        cursor.execute("PRAGMA table_info(users)")
        users_columns = [row[1] for row in cursor.fetchall()]
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        # conn.close() # 优化：使用共享连接
        
        return {
            "success": True, 
            "tables": tables,
            "users_columns": users_columns,
            "added_columns": missing_columns,
            "message": "数据库表结构已修复"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/debug/create-admin")
async def create_admin(
    email: str = Query(default="zhwffy@hotmail.com"),
    name: str = Query(default="Admin"),
    password: str = Query(default="Hiller")
):
    """创建管理员账号"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        
        # 检查用户是否存在
        cursor.execute("SELECT id FROM users WHERE email = ?", (email,))
        if cursor.fetchone():
            # 更新为管理员
            cursor.execute("UPDATE users SET is_admin = 1 WHERE email = ?", (email,))
            conn.commit()
            # conn.close() # 优化：使用共享连接
            return {"success": True, "message": "已将用户设置为管理员", "email": email}
        
        # 创建新管理员
        password_hash = hash_password(password)
        cursor.execute('''INSERT INTO users 
            (email, name, password_hash, is_admin, is_paid, created_at)
            VALUES (?, ?, ?, 1, 1, ?)''', 
            (email, name, password_hash, datetime.now().isoformat()))
        conn.commit()
        user_id = cursor.lastrowid
        # conn.close() # 优化：使用共享连接
        
        # 生成token
        token = create_token(user_id, email)
        
        return {
            "success": True,
            "message": "管理员账号已创建",
            "user": {
                "id": user_id,
                "email": email,
                "name": name,
                "is_admin": True
            },
            "token": token
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/debug/test-email")
async def test_email(request: SendCodeRequest):
    """测试邮件发送"""
    try:
        code = generate_code(EMAIL_CONFIG["code_length"])
        result = send_verification_code(request.email, code)
        return {
            "success": result["success"],
            "message": result["message"],
            "test_mode": EMAIL_CONFIG["test_mode"],
            "smtp_server": EMAIL_CONFIG["smtp_server"],
            "smtp_user": EMAIL_CONFIG["smtp_user"],
            "error": result.get("error") if not result["success"] else None
        }
    except Exception as e:
        return {"success": False, "error": str(e), "error_type": type(e).__name__}

@app.get("/debug/users")
async def debug_users():
    """查看所有用户"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        cursor.execute("SELECT id, email, name, password_hash, created_at FROM users")
        users = []
        for row in cursor.fetchall():
            users.append({
                "id": row["id"],
                "email": row["email"],
                "name": row["name"],
                "password_hash": row["password_hash"],
                "created_at": row["created_at"]
            })
        # conn.close() # 优化：使用共享连接
        return {"users": users, "count": len(users)}
    except Exception as e:
        return {"error": str(e)}

@app.post("/debug/check-password")
async def debug_check_password(email: str, password: str):
    """检查密码匹配"""
    try:
        conn = get_user_db()
        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()
        # conn.close() # 优化：使用共享连接
        
        if not user:
            return {"error": "用户不存在", "email": email}
        
        stored_hash = user["password_hash"]
        input_hash = hash_password(password)
        
        return {
            "email": email,
            "stored_hash": stored_hash,
            "input_hash": input_hash,
            "match": stored_hash == input_hash,
            "password": password
        }
    except Exception as e:
        return {"error": str(e)}

# ========== 管理员API ==========

async def get_admin_user(authorization: Optional[str] = Header(None)):
    """验证管理员权限"""
    if not authorization:
        raise HTTPException(status_code=401, detail="未提供认证token")
    
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    payload = verify_token(token)
    
    if not payload:
        raise HTTPException(status_code=401, detail="token无效或已过期")
    
    conn = get_user_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (payload["user_id"],))
    user = cursor.fetchone()
    # conn.close() # 优化：使用共享连接
    
    if not user:
        raise HTTPException(status_code=401, detail="用户不存在")
    
    if not user["is_admin"]:
        raise HTTPException(status_code=403, detail="需要管理员权限")
    
    return dict(user)

@app.get("/api/v1/admin/stats")
async def admin_stats(admin: dict = Depends(get_admin_user)):
    """获取系统统计（带缓存优化）"""
    
    # 检查缓存
    cached = get_cached_stats()
    if cached:
        return {"success": True, "data": cached, "cached": True}
    
    conn = get_user_db()
    cursor = conn.cursor()
    
    # 用户统计
    cursor.execute("SELECT COUNT(*) as count FROM users")
    total_users = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE is_paid = 1")
    paid_users = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM users WHERE DATE(created_at) = DATE('now')")
    new_users_today = cursor.fetchone()["count"]
    
    # 使用统计
    cursor.execute("SELECT COUNT(*) as count FROM usage")
    total_usage = cursor.fetchone()["count"]
    
    cursor.execute("SELECT COUNT(*) as count FROM usage WHERE DATE(created_at) = DATE('now')")
    usage_today = cursor.fetchone()["count"]
    
    # 简历统计
    cursor.execute("SELECT COUNT(*) as count FROM resumes")
    total_resumes = cursor.fetchone()["count"]
    
    # 等级分布统计
    cursor.execute("""
        SELECT 
            CASE 
                WHEN is_admin = 1 THEN 'admin'
                WHEN is_paid = 1 AND user_level = 'vip' THEN 'vip'
                WHEN is_paid = 1 AND user_level = 'pro' THEN 'pro'
                WHEN is_paid = 1 AND user_level = 'basic' THEN 'basic'
                ELSE 'free'
            END as level,
            COUNT(*) as count
        FROM users
        GROUP BY level
    """)
    level_distribution = {row["level"]: row["count"] for row in cursor.fetchall()}
    
    # VIP用户数（付费用户）
    vip_users = paid_users
    
    stats_data = {
        "total_users": total_users,
        "today_usage": usage_today,
        "total_usage": total_usage,
        "vip_users": vip_users,
        "level_distribution": level_distribution,
        "email_config": get_email_config()
    }
    
    # 更新缓存
    set_cached_stats(stats_data)
    
    return {
        "success": True,
        "data": stats_data,
        "cached": False
    }

@app.get("/api/v1/admin/users")
async def admin_list_users(admin: dict = Depends(get_admin_user)):
    """获取用户列表"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            u.id, u.email, u.name, u.is_paid, u.is_admin, u.created_at, u.last_login,
            (SELECT COUNT(*) FROM resumes WHERE user_id = u.id) as resume_count,
            (SELECT COUNT(*) FROM usage WHERE user_id = u.id) as usage_count
        FROM users u
        ORDER BY u.created_at DESC
    """)
    
    users = []
    for row in cursor.fetchall():
        users.append({
            "id": row["id"],
            "email": row["email"],
            "name": row["name"],
            "is_paid": bool(row["is_paid"]),
            "is_admin": bool(row["is_admin"]),
            "created_at": row["created_at"],
            "last_login": row["last_login"],
            "resume_count": row["resume_count"],
            "usage_count": row["usage_count"]
        })
    
    # conn.close() # 优化：使用共享连接
    return {"success": True, "data": users}

class UpdateUserRequest(BaseModel):
    is_paid: Optional[bool] = None
    is_admin: Optional[bool] = None

@app.put("/api/v1/admin/users/{user_id}")
async def admin_update_user(user_id: int, request: UpdateUserRequest, admin: dict = Depends(get_admin_user)):
    """更新用户状态"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    updates = []
    params = []
    
    if request.is_paid is not None:
        updates.append("is_paid = ?")
        params.append(1 if request.is_paid else 0)
    
    if request.is_admin is not None:
        updates.append("is_admin = ?")
        params.append(1 if request.is_admin else 0)
    
    if not updates:
        # conn.close() # 优化：使用共享连接
        raise HTTPException(status_code=400, detail="没有要更新的字段")
    
    params.append(user_id)
    cursor.execute(f"UPDATE users SET {', '.join(updates)} WHERE id = ?", params)
    conn.commit()
    # conn.close() # 优化：使用共享连接
    
    return {"success": True, "message": "用户已更新"}

@app.get("/api/v1/admin/usage")
async def admin_usage_logs(admin: dict = Depends(get_admin_user)):
    """获取使用日志"""
    conn = get_user_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT 
            u.id, u.user_id, u.action, u.industry, u.created_at,
            us.email, us.name as user_name
        FROM usage u
        LEFT JOIN users us ON u.user_id = us.id
        ORDER BY u.created_at DESC
        LIMIT 100
    """)
    
    logs = []
    for row in cursor.fetchall():
        logs.append({
            "id": row["id"],
            "user_id": row["user_id"],
            "email": row["email"] or "-",
            "user_name": row["user_name"] or "-",
            "action": row["action"],
            "industry": row["industry"],
            "created_at": row["created_at"]
        })
    
    # conn.close() # 优化：使用共享连接
    return {"success": True, "data": logs}

class TestAIRequest(BaseModel):
    provider: str = "dashscope"
    model: str = "qwen-turbo"
    api_key: Optional[str] = None
    base_url: Optional[str] = None

@app.post("/api/v1/admin/test-ai")
async def admin_test_ai(request: TestAIRequest, admin: dict = Depends(get_admin_user)):
    """测试AI连接"""
    import httpx
    
    # 使用请求中的配置或默认配置
    api_key = request.api_key or DASHSCOPE_API_KEY
    base_url = request.base_url or DASHSCOPE_BASE_URL
    model = request.model or DASHSCOPE_MODEL
    
    if not api_key:
        return {"success": False, "error": "未配置API Key"}
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{base_url}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "User-Agent": "Mozilla/5.0 ResumeAI/1.0"  # Coding API需要User-Agent
                },
                json={
                    "model": model,
                    "messages": [
                        {"role": "user", "content": "请回复：AI连接测试成功"}
                    ],
                    "max_tokens": 50
                }
            )
            
            if response.status_code == 200:
                data = response.json()
                reply = data.get("choices", [{}])[0].get("message", {}).get("content", "")
                return {"success": True, "response": reply, "model": model}
            else:
                error_msg = response.text[:200]
                return {"success": False, "error": f"API错误: {response.status_code} - {error_msg}"}
    except Exception as e:
        return {"success": False, "error": f"连接失败: {str(e)}"}

@app.get("/api/v1/admin/config")
async def admin_get_config(admin: dict = Depends(get_admin_user)):
    """获取系统配置"""
    config = {
        "email": {
            "smtp_server": EMAIL_CONFIG.get("smtp_server", ""),
            "smtp_port": EMAIL_CONFIG.get("smtp_port", 465),
            "smtp_user": EMAIL_CONFIG.get("smtp_user", ""),
            "configured": EMAIL_CONFIG.get("smtp_configured", False),
            "test_mode": EMAIL_CONFIG.get("test_mode", True)
        },
        "ai": {
            "provider": "coding",
            "model": DASHSCOPE_MODEL,
            "base_url": DASHSCOPE_BASE_URL,
            "configured": bool(DASHSCOPE_API_KEY)
        }
    }
    return {"success": True, "data": config}

if __name__ == "__main__":
    import uvicorn
    print(f"🚀 ResumeAI Backend Starting...")
    print(f"Keywords loaded: {len(KEYWORDS_DB)} industries")
    print(f"Auth enabled: True")
    print(f"Free limit: {FREE_LIMIT} times/day")
    uvicorn.run(app, host="0.0.0.0", port=8001)

# Vercel Serverless Handler
try:
    from mangum import Mangum
    handler = Mangum(app, lifespan="off")
except ImportError:
    # 如果mangum未安装，提供一个基本的handler
    def handler(event, context):
        return {
            "statusCode": 500,
            "body": "Mangum not installed",
            "headers": {"Content-Type": "text/plain"}
        }